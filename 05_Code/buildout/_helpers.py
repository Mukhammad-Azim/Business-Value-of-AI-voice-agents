"""
Shared helpers for the thesis build-out scripts.

All scripts in 08_Code/buildout/ open the same `04_Drafts/Main_Draft.docx`,
mutate it in place, and save. Each script is idempotent: running it twice
must not produce duplicated content. Idempotency is enforced with sentinel
markers stored as italic / hidden runs at the start of inserted blocks
("[BUILDOUT:lit_table_1]", "[BUILDOUT:ch6]", etc.).

Helpers exposed:

* ``insert_paragraph_before(anchor, text, style)``
* ``insert_picture_before(anchor, image_path, width_in)``
* ``insert_table_before(anchor, n_rows, n_cols, style_name)``
* ``insert_paragraph_after(anchor, text, style)``
* ``has_marker(doc, marker)``
* ``add_marker(paragraph, marker)``
"""

from __future__ import annotations

import os

import docx
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPO_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DRAFT_PATH = os.path.join(REPO_ROOT, "04_Drafts", "Main_Draft.docx")
OUTPUTS = os.path.join(REPO_ROOT, "outputs")
FIG_FIN = os.path.join(OUTPUTS, "figures", "financial")
FIG_NLP = os.path.join(OUTPUTS, "figures", "NLP")
FIG_DESIGN = os.path.join(OUTPUTS, "figures", "design_pattern")


def open_draft():
    return docx.Document(DRAFT_PATH)


def save_draft(doc):
    doc.save(DRAFT_PATH)


def insert_paragraph_before(anchor_para, text="", style_name=None):
    """Insert a new paragraph (with given text + style) directly before the
    anchor paragraph. Returns the new Paragraph object."""
    doc = anchor_para.part.document
    tmp = doc.add_paragraph()
    if style_name:
        try:
            tmp.style = doc.styles[style_name]
        except KeyError:
            pass
    if text:
        tmp.add_run(text)
    anchor_para._p.addprevious(tmp._p)
    return tmp


def insert_paragraph_after(anchor_para, text="", style_name=None):
    """Insert a paragraph immediately after the anchor."""
    doc = anchor_para.part.document
    tmp = doc.add_paragraph()
    if style_name:
        try:
            tmp.style = doc.styles[style_name]
        except KeyError:
            pass
    if text:
        tmp.add_run(text)
    anchor_para._p.addnext(tmp._p)
    return tmp


def insert_picture_before(anchor_para, image_path, width_in=6.0):
    doc = anchor_para.part.document
    tmp = doc.add_paragraph()
    tmp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tmp.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    anchor_para._p.addprevious(tmp._p)
    return tmp


def insert_table_before(anchor_para, n_rows, n_cols, style_name=None):
    """Insert a fresh table before the anchor paragraph.
    Returns the table object."""
    doc = anchor_para.part.document
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    if style_name:
        try:
            tbl.style = doc.styles[style_name]
        except KeyError:
            pass
    anchor_para._p.addprevious(tbl._tbl)
    return tbl


def has_marker(doc, marker):
    """True if any paragraph contains the marker text (used as idempotency
    sentinel). Markers are written as italic small runs prefixed with
    'BUILDOUT::'."""
    needle = f"BUILDOUT::{marker}"
    for p in doc.paragraphs:
        if needle in p.text:
            return True
    return False


def add_marker(paragraph, marker):
    """Append an invisible-ish marker run to a paragraph (small, italic)."""
    needle = f"BUILDOUT::{marker}"
    run = paragraph.add_run(" " + needle)
    run.italic = True
    run.font.size = Pt(1)
    return run


def find_paragraph_index_starting(doc, prefix):
    """Return the first paragraph whose stripped text starts with `prefix`."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i, p
    return -1, None


def find_paragraph_with_text(doc, needle):
    """Return the first paragraph whose stripped text contains needle."""
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return i, p
    return -1, None


def set_table_cell(cell, text, bold=False, italic=False, font_size=None):
    """Replace cell content with formatted text."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    return p


def styled_para_runs(paragraph, parts):
    """Append styled runs to a paragraph. `parts` is a list of (text, kwargs)
    where kwargs accept bold=, italic=, underline=."""
    for text, opts in parts:
        run = paragraph.add_run(text)
        for k, v in opts.items():
            setattr(run, k, v)


def delete_table(tbl):
    """Delete a python-docx Table from its parent."""
    el = tbl._tbl
    el.getparent().remove(el)


def insert_paragraph_at_end(doc, text="", style_name=None):
    """Append a new paragraph at the end of the document."""
    p = doc.add_paragraph()
    if style_name:
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            pass
    if text:
        p.add_run(text)
    return p


def insert_picture_at_end(doc, image_path, width_in=6.0):
    """Append a new paragraph with an embedded picture at the end."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    return p


def insert_table_at_end(doc, n_rows, n_cols, style_name=None):
    """Append a new table at the end of the document."""
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    if style_name:
        try:
            tbl.style = doc.styles[style_name]
        except KeyError:
            pass
    return tbl


def list_tables_in_doc(doc):
    """Return list of (table_index, body_position) tuples in document order."""
    body = doc.element.body
    pos = 0
    table_idx = 0
    out = []
    for child in body.iterchildren():
        if child.tag.endswith("}tbl"):
            out.append((table_idx, pos))
            table_idx += 1
        pos += 1
    return out
