"""
11_chapter9_synthesise.py
=========================

Rewrite Chapter 9 so it reads as a synthetic conclusion rather than a
restatement of Chapter 5. The previous draft repeated full statistical
detail (p-values, Cohen's d, sample sizes, Bonferroni and Welch
qualifiers, four-decimal sentiment values) verbatim in §9.2. Per the
supervisor's instruction these belong to Chapter 5 and the appendices;
Chapter 9 should retain only headline values that help answer RQ1 and
RQ2.

Specific edits implemented here:

* §9.1 (RQ1): keep the four-prerequisite synthesis, add a single
  cross-reference to Chapter 4 / Chapter 6 so the reader knows where the
  detailed evidence lives. Wording remains non-causal.

* §9.2 (RQ2): replace the long paragraph with a tighter synthesis that
  retains four headline values (the expense-ratio improvement
  4.79 -> 2.01, the revenue-growth versus expense-growth contrast, the
  revenue-per-employee movement from the 2024 trough back towards the
  pre-period level, and the qualitative direction of the
  NPS-versus-Trustpilot divergence) and cross-references the relevant
  Chapter 5 sections for the full numeric detail. p-values, Cohen's d,
  exact sample sizes, Bonferroni and Welch qualifiers, four-decimal
  sentiment values, and full keyword-shift magnitudes are removed from
  Chapter 9.

* §9.3.1 (Empirical Contribution): fix the dangling "the thesis is the
  an early" wording.

* New §9.4 (Limitations): brief paragraph that acknowledges the main
  limitations and cross-references Chapter 8 for the full treatment.

* §9.4 -> §9.5 Future Research, §9.5 -> §9.6 Closing Statement
  (renumbered to preserve the new §9.4 limitations section).

The script is idempotent: if the new wording is already in place the
script makes no change.
"""

from __future__ import annotations

import sys

sys.path.insert(0, ".")
from _helpers import open_draft, save_draft, insert_paragraph_after  # noqa: E402


# ---------------------------------------------------------------------------
# Replacement targets (full paragraph -> full paragraph)
# ---------------------------------------------------------------------------

OLD_91 = (
    "RQ1 asked how voice LLM agents have been integrated into Better.com\u2019s "
    "mortgage origination processes and what this integration reveals about "
    "the technical and organisational prerequisites for deploying conversational "
    "AI in a regulated financial-services environment. The case study supports "
    "a four-part answer. First, the integration is anchored in a process-aware "
    "operating platform (Tinman) that exposes structured workflow state, "
    "document state, eligibility state, and timeline events through "
    "permissioned interfaces. Second, the voice agent (Betsy) is scoped to the "
    "coordination layer of the workflow \u2014 lead engagement, scheduling, "
    "status updates, document collection, post-decision communication \u2014 "
    "and is not scoped to underwriting, credit decisioning, or the disclosure "
    "of material loan terms. Third, the voice infrastructure (ElevenLabs "
    "Agents) supplies low-latency speech-to-text, response generation, and "
    "text-to-speech, with the operational latency target consistent with the "
    "latency thresholds documented in the human-computer interaction literature "
    "(Maslych et al., 2025). Fourth, the regulated boundary is operationalised "
    "through an auditable human-AI collaboration model in which licensed loan "
    "consultants handle regulated steps and the agent\u2013human transitions "
    "are timestamped and audit-logged. Read together, these four prerequisites "
    "are the technical and organisational conditions the case identifies for "
    "deploying conversational AI in a regulated financial-services environment."
)

NEW_91 = (
    "RQ1 asked how voice LLM agents have been integrated into Better.com\u2019s "
    "mortgage origination processes and what this integration reveals about "
    "the technical and organisational prerequisites for deploying "
    "conversational AI in a regulated financial-services environment. The case "
    "supports a four-part answer, formalised as the design pattern presented "
    "in Chapter 6. First, the integration is anchored in a process-aware "
    "operating platform (Tinman) that exposes structured workflow, document, "
    "eligibility, and timeline state through permissioned interfaces "
    "(see \u00a74.5). Second, the voice agent (Betsy) is scoped to the "
    "coordination layer of the workflow \u2014 lead engagement, scheduling, "
    "status updates, document collection, and post-decision communication "
    "\u2014 and is not scoped to underwriting, credit decisioning, or the "
    "disclosure of material loan terms. Third, the voice infrastructure "
    "(ElevenLabs Agents) supplies low-latency speech-to-text, response "
    "generation, and text-to-speech, with the operational latency target "
    "consistent with the human-computer interaction literature (Maslych "
    "et al., 2025). Fourth, the regulated boundary is operationalised "
    "through an auditable human\u2013AI collaboration model in which "
    "licensed loan consultants handle regulated steps and agent\u2013human "
    "transitions are timestamped and audit-logged. Read together, these "
    "four prerequisites are the technical and organisational conditions the "
    "case identifies for deploying conversational AI in regulated "
    "financial-services workflows."
)

# §9.2 — the heavy stats restatement that needs to go.
NEW_92 = (
    "RQ2 asked what measurable business value can be associated with "
    "Better.com\u2019s voice-AI deployment across operational efficiency, "
    "productivity, customer experience, and financial performance, and what "
    "factors mediate the realisation of that value. Synthesising the "
    "triangulated evidence reported in Chapter 5, the answer has three parts. "
    "First, the post-deployment trajectory is consistent with operational "
    "leverage: the expense ratio improved from 4.79 in 2023 to 2.01 in 2025, "
    "revenue growth in 2024\u20132025 substantially outpaced expense growth, "
    "and revenue-per-employee recovered from its 2024 trough back towards the "
    "pre-period level (see \u00a75.2 for the full numeric detail and "
    "\u00a75.6 for the macro-environment caveats). Second, the customer-"
    "experience signal is divergent across instruments: management-reported "
    "NPS and CSAT improved over the same window, while the independent "
    "Trustpilot corpus shows a post-deployment decline in mean star rating "
    "and in VADER sentiment, alongside a contraction of Cost/Rates discourse "
    "as the only keyword-prevalence shift surviving multiple-comparison "
    "correction (see \u00a75.3 and \u00a75.4 for the full statistical "
    "treatment, including Welch tests, effect sizes, and the BERTopic and "
    "keyword-prevalence diagnostics). Third, the value-realisation factors "
    "are the four prerequisites of the design pattern: the process-aware "
    "platform, the coordination-layer agent scope, the low-latency voice "
    "infrastructure, and the auditable human-escalation boundary. The "
    "thesis maintains that the financial trajectory is consistent with "
    "AI-enabled operational leverage but cannot be attributed to Betsy in "
    "isolation; the contributions of macroeconomic recovery, restructuring, "
    "and Tinman-platform maturity are not separable from the voice-AI "
    "contribution in the case data, and the thesis cannot isolate Betsy\u2019s "
    "independent causal contribution."
)

OLD_931 = (
    "The thesis is the an early published case study, to the author\u2019s "
    "knowledge, to triangulate audited SEC-filed financial trajectory, "
    "independent customer-discourse mining, and qualitative process tracing "
    "of integration architecture for a voice LLM agent in U.S. mortgage "
    "origination. The single-case design is deliberately revelatory rather "
    "than statistical, and the thesis records the trajectory and the "
    "discourse signal in a form that is reproducible from public sources."
)

NEW_931 = (
    "The thesis is, to the author\u2019s knowledge, an early published case "
    "study to triangulate the audited SEC-filed financial trajectory, "
    "independent customer-discourse mining, and qualitative process tracing "
    "of integration architecture for a voice LLM agent in U.S. mortgage "
    "origination. The single-case design is deliberately revelatory rather "
    "than statistical, and the thesis records the trajectory and the "
    "discourse signal in a form that is reproducible from public sources."
)

# New §9.4 Limitations heading + paragraph (inserted before existing 9.4).
NEW_94_BODY = (
    "Chapter 8 sets out the full limitations of the thesis; the conclusion "
    "summarises only the four most consequential. The single-case "
    "explanatory design constrains statistical generalisability and "
    "precludes a causal claim about Betsy\u2019s independent contribution. "
    "The empirical window contains a substantial macroeconomic and product-"
    "mix recovery in U.S. mortgage origination, which is not separable from "
    "the voice-AI deployment in the case data. The Trustpilot corpus is a "
    "self-selected, non-random sample of consumers and is not "
    "representative of the customer base, even though its post-deployment "
    "decline survived multiple-comparison correction on the Cost/Rates "
    "topic. Finally, several integration-architecture descriptions rely on "
    "vendor- and management-reported materials whose source-reliability "
    "tier is explicitly flagged in \u00a73.7 and Appendix F, and which "
    "should not be read as audited evidence."
)


def replace_paragraph_text(paragraph, new_text: str) -> bool:
    """Replace paragraph text in place. Preserves the first run's
    formatting; clears all subsequent runs. Returns True if changed."""
    runs = paragraph.runs
    current = "".join(r.text for r in runs)
    if current.strip() == new_text.strip():
        return False
    if not runs:
        # Fall back to direct text assignment via the underlying XML.
        paragraph.text = new_text
        return True
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def find_paragraph_starting_with(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def main():
    doc = open_draft()
    changed = []

    # 1. Rewrite §9.1, §9.2, §9.3.1.
    p91 = find_paragraph_starting_with(doc, "RQ1 asked how voice LLM agents")
    if p91 is not None and replace_paragraph_text(p91, NEW_91):
        changed.append("9.1 RQ1 paragraph rewritten")

    p92 = find_paragraph_starting_with(doc, "RQ2 asked what measurable business value")
    if p92 is not None and replace_paragraph_text(p92, NEW_92):
        changed.append("9.2 RQ2 paragraph rewritten (verbose stats removed)")

    p931 = find_paragraph_starting_with(doc, "The thesis is the an early")
    if p931 is not None and replace_paragraph_text(p931, NEW_931):
        changed.append("9.3.1 typo fixed")

    # 2. Insert §9.4 Limitations *before* the current §9.4 Future Research
    #    heading and renumber the following two H2 headings.
    p_future = None
    p_closing = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip().startswith("9.4 Future Research"):
            p_future = p
        if p.style.name == "Heading 2" and p.text.strip().startswith("9.5 Closing Statement"):
            p_closing = p

    needs_insert = True
    # Detect idempotency: if §9.4 Limitations heading already exists, skip insertion.
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip().startswith("9.4 Limitations"):
            needs_insert = False
            break

    if needs_insert and p_future is not None:
        # Renumber the existing §9.4 -> §9.5 and §9.5 -> §9.6 first.
        replace_paragraph_text(p_future, "9.5 Future Research")
        if p_closing is not None:
            replace_paragraph_text(p_closing, "9.6 Closing Statement")

        # Insert new §9.4 heading and body BEFORE the (now renumbered)
        # §9.5 Future Research paragraph.
        # Strategy: insert a Heading 2 paragraph followed by a Normal
        # paragraph, both before the existing §9.5.
        # python-docx doesn't have insert_paragraph_before on Document,
        # so we use the paragraph element's lxml addprevious() method via
        # a helper.
        from docx.oxml.ns import qn
        from copy import deepcopy

        # 1) Build a Heading 2 paragraph by deep-copying the current 9.5
        #    heading XML and rewriting its text.
        h2_xml = deepcopy(p_future._p)
        # Strip every <w:r> child and replace with a single run carrying
        # the new heading text.
        for child in list(h2_xml):
            if child.tag == qn("w:r"):
                h2_xml.remove(child)
        # Now add a single run with the heading text by reusing the
        # existing run-level rPr (if any) from p_future's first run.
        from docx.oxml import OxmlElement
        run_el = OxmlElement("w:r")
        if p_future.runs:
            existing_rpr = p_future.runs[0]._r.find(qn("w:rPr"))
            if existing_rpr is not None:
                run_el.append(deepcopy(existing_rpr))
        text_el = OxmlElement("w:t")
        text_el.text = "9.4 Limitations"
        text_el.set(qn("xml:space"), "preserve")
        run_el.append(text_el)
        h2_xml.append(run_el)

        # 2) Build a Normal paragraph for the body text by deep-copying
        #    the §9.3.3 Practical body paragraph (we know it's a Normal
        #    style) and rewriting its content.
        p93_body = find_paragraph_starting_with(doc, "For lender executives, the four-prerequisite")
        if p93_body is None:
            # Fallback to any Normal-style paragraph adjacent to §9.4.
            p93_body = doc.paragraphs[doc.paragraphs.index(p_future) - 1]
        body_xml = deepcopy(p93_body._p)
        for child in list(body_xml):
            if child.tag == qn("w:r"):
                body_xml.remove(child)
        body_run_el = OxmlElement("w:r")
        if p93_body.runs:
            existing_rpr = p93_body.runs[0]._r.find(qn("w:rPr"))
            if existing_rpr is not None:
                body_run_el.append(deepcopy(existing_rpr))
        body_text_el = OxmlElement("w:t")
        body_text_el.text = NEW_94_BODY
        body_text_el.set(qn("xml:space"), "preserve")
        body_run_el.append(body_text_el)
        body_xml.append(body_run_el)

        # Insert the heading and body XML elements before the §9.5
        # Future Research heading.
        p_future._p.addprevious(h2_xml)
        p_future._p.addprevious(body_xml)
        changed.append("9.4 Limitations subsection inserted (before §9.5)")
        changed.append("9.4 Future Research renumbered to 9.5")
        if p_closing is not None:
            changed.append("9.5 Closing Statement renumbered to 9.6")

    save_draft(doc)

    print("=" * 60)
    print("Chapter 9 synthetic-conclusion pass")
    print("=" * 60)
    if not changed:
        print("No changes — Chapter 9 already in revised form.")
    else:
        for c in changed:
            print(f"  - {c}")


if __name__ == "__main__":
    main()
