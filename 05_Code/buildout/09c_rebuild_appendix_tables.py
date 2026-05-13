"""
09c_rebuild_appendix_tables.py
==============================

Reconciles Appendices A, B, C, and D with the actual project data files
(supervisor blocking issues A3, A4, A5, A6, A7).

Authoritative sources used to reconstruct the tables:

* ``Data/financial_dataset_annual.csv`` (13 rows: 5 Better 2021-2025,
  8 Rocket 2018-2025).
* ``08_Code/Data Extraction code/financial_dataset_documentation.md``
  (filing-by-filing provenance and the SRC-rule note on Better 2021).
* ``03_Case_Study_Data/Rocket Mortgage/Annual Reports/`` and
  ``03_Case_Study_Data/Better/Annual Reports/`` (filing PDFs on disk).
* ``Data/NLP/Better_Reviews.csv`` and ``08_Code/cleaning_reviews.ipynb``
  (the executed Trustpilot cleaning sequence).
* ``08_Code/better_trustpilot_keyword_analysis.csv`` and
  ``08_Code/better_trustpilot_with_topics.csv`` (the seven keyword groups
  and the BERTopic topic distribution actually computed).

Idempotency: the script overwrites the table contents in place, so re-running
it on a clean draft simply rewrites the same data. No sentinel marker is
required.
"""

from __future__ import annotations

import os
import sys

import pandas as pd

sys.path.insert(0, ".")
from _helpers import open_draft, save_draft, REPO_ROOT  # noqa: E402

CSV_FIN = os.path.join(REPO_ROOT, "Data", "financial_dataset_annual.csv")
CSV_KW = os.path.join(REPO_ROOT, "08_Code", "better_trustpilot_keyword_analysis.csv")
CSV_TOPICS = os.path.join(REPO_ROOT, "08_Code", "better_trustpilot_with_topics.csv")


# ---------------------------------------------------------------------------
# Table A.2 — Annual financial observations (13 rows)
# ---------------------------------------------------------------------------

def build_table_a2_rows():
    df = pd.read_csv(CSV_FIN)
    rows = []
    for _, r in df.iterrows():
        firm = "Better Home & Finance Holding Company" if r["Company"] == "Better" else "Rocket Companies, Inc."
        net_loss = r["Net_Income"]  # negative for loss
        rows.append({
            "Firm": firm,
            "FY": int(r["Year"]),
            "Revenue (USDm)": f"{r['Revenue']:,.1f}",
            "Total expenses (USDm)": f"{r['Total_Expenses']:,.1f}",
            "Net income / (loss) (USDm)": (
                f"({abs(net_loss):,.1f})" if net_loss < 0 else f"{net_loss:,.1f}"
            ),
            "Headcount": f"{int(r['Employees']):,}",
            "Expense ratio": f"{r['Expense_Ratio']:.3f}",
            "RPE (USDm)": f"{r['Revenue_per_Employee']:.3f}",
        })
    return rows


def rewrite_table_a2(doc):
    """Locate Table A.2 by its header row and rewrite all data rows.

    The current Table A.2 has 9 rows (1 header + 8 obs) but the dataset has
    13 obs, so we recreate the table at the same XML position with 14 rows
    (1 header + 13 obs). To keep formatting consistent with python-docx, we
    delete excess rows or add new rows as needed.
    """
    target = None
    for t in doc.tables:
        if not t.rows:
            continue
        h = [c.text.strip() for c in t.rows[0].cells]
        if h[:3] == ["Company", "FY", "Revenue (USDm)"] or (
            len(h) >= 4 and h[1] == "FY" and h[2].startswith("Revenue")
        ):
            target = t
            break
    if target is None:
        print("[A.2] table not found")
        return False

    rows = build_table_a2_rows()
    # New header (8 columns, restated for clarity).
    new_header = [
        "Firm", "FY", "Revenue (USDm)", "Total expenses (USDm)",
        "Net income / (loss) (USDm)", "Headcount",
        "Expense ratio", "RPE (USDm)",
    ]
    # Ensure number of columns matches.
    if len(target.rows[0].cells) != len(new_header):
        print(f"[A.2] column count mismatch ({len(target.rows[0].cells)} vs {len(new_header)})")
        return False

    # Trim or extend rows to match.
    while len(target.rows) > 1:
        # Remove the second row repeatedly until only the header remains.
        target._tbl.remove(target.rows[1]._tr)

    # Update header cells.
    for i, txt in enumerate(new_header):
        cell = target.rows[0].cells[i]
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        cell.paragraphs[0].add_run(txt).bold = True

    # Append data rows.
    for row in rows:
        new_row = target.add_row().cells
        for i, key in enumerate(new_header):
            cell = new_row[i]
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            cell.paragraphs[0].add_run(str(row[key]))

    print(f"[A.2] rewritten with {len(rows)} data rows")
    return True


# ---------------------------------------------------------------------------
# Table A.1 — Variable definitions: tighten metadata
# ---------------------------------------------------------------------------

def rewrite_table_a1(doc):
    """Replace Table A.1 with a clean variable-definitions table whose 'Computed
    from' column references Table A.2 line items rather than implying audited
    deterministic ratios where some are author-constructed."""
    target = None
    for t in doc.tables:
        if not t.rows:
            continue
        h = [c.text.strip() for c in t.rows[0].cells]
        if h[:3] == ["#", "Variable", "Definition"]:
            target = t
            break
    if target is None:
        print("[A.1] table not found")
        return False

    new_rows = [
        ("#", "Variable", "Definition", "Unit", "Computed from"),
        ("1", "Revenue", "Total net revenue from the audited Consolidated Statement of Operations.", "USD millions", "10-K / S-1; line item: Total revenue, net."),
        ("2", "Total expenses", "Total operating expenses from the audited Consolidated Statement of Operations.", "USD millions", "10-K / S-1; line item: Total expenses (sum of compensation, marketing, technology, G&A, depreciation)."),
        ("3", "Net income / (loss)", "Net income or net loss attributable to common shareholders.", "USD millions", "10-K / S-1; line item: Net income / (loss)."),
        ("4", "Operating cash flow", "Net cash provided by / (used in) operating activities.", "USD millions", "10-K / S-1; Consolidated Statement of Cash Flows."),
        ("5", "Total assets", "Total assets, audited balance-sheet aggregate.", "USD millions", "10-K / S-1; Consolidated Balance Sheet."),
        ("6", "Total equity", "Common equity attributable to the parent.", "USD millions", "10-K / S-1; Consolidated Balance Sheet."),
        ("7", "Headcount", "Full-time-equivalent employees as disclosed in Item 1 (Human Capital) of the 10-K.", "FTE", "10-K Item 1; Better S-1 narrative for FY 2021."),
        ("8", "Expense ratio", "Ratio of total expenses to total revenue.", "ratio", "Computed: Total expenses / Revenue."),
        ("9", "Revenue per employee (RPE)", "Total revenue divided by headcount.", "USD millions / FTE", "Computed: Revenue / Headcount."),
        ("10", "Cost-to-Originate (CTO)", "Author-constructed proxy: total expenses divided by funded loans, used only as a directional indicator alongside the Mortgage Bankers Association industry benchmark.", "USD per loan", "Author-constructed (not an audited Better disclosure)."),
    ]

    while len(target.rows) > 1:
        target._tbl.remove(target.rows[1]._tr)

    # Update header row.
    for i, txt in enumerate(new_rows[0]):
        cell = target.rows[0].cells[i]
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        cell.paragraphs[0].add_run(txt).bold = True

    # Append data rows.
    for row in new_rows[1:]:
        new_row = target.add_row().cells
        for i, txt in enumerate(row):
            cell = new_row[i]
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            cell.paragraphs[0].add_run(str(txt))

    print("[A.1] variable-definitions table rewritten")
    return True


# ---------------------------------------------------------------------------
# Table B.1 — SEC filing provenance (full Rocket 2020-2025 + Better)
# ---------------------------------------------------------------------------

def rewrite_table_b1(doc):
    target = None
    for t in doc.tables:
        if not t.rows:
            continue
        h = [c.text.strip() for c in t.rows[0].cells]
        if h[:3] == ["#", "Filer", "Form"]:
            target = t
            break
    if target is None:
        print("[B.1] table not found")
        return False

    new_rows = [
        ("#", "Filer", "Form", "Reporting period", "Use in this thesis", "Repository path"),
        ("1", "Better HoldCo, Inc.", "S-1 / S-1A (registration statement)", "FY 2020-2022 historical financials", "Source for FY 2021 audited income-statement and headcount values used in Table A.2 (Better is a Smaller Reporting Company; the 2023 10-K reports only two years).", "03_Case_Study_Data/Better/Annual Reports/S1.pdf"),
        ("2", "Better Home & Finance Holding Company", "10-K", "FY 2023 (filed 2024)", "Source for FY 2022 and FY 2023 audited income-statement, balance-sheet, and headcount values used in Table A.2.", "03_Case_Study_Data/Better/Annual Reports/2023.pdf"),
        ("3", "Better Home & Finance Holding Company", "10-K", "FY 2024 (filed 2025)", "Source for FY 2024 audited income-statement, balance-sheet, and headcount values used in Table A.2; also discloses the October 2024 Betsy launch in the operational narrative.", "03_Case_Study_Data/Better/Annual Reports/2024.pdf"),
        ("4", "Better Home & Finance Holding Company", "10-K", "FY 2025 (filed 2026)", "Source for FY 2025 audited income-statement, balance-sheet, and headcount values used in Table A.2.", "03_Case_Study_Data/Better/Annual Reports/2025.pdf"),
        ("5", "Rocket Companies, Inc.", "10-K", "FY 2020 (filed 2021)", "Source for FY 2020 audited values used in Table A.2; this filing also contains the FY 2018 and FY 2019 audited comparative tables under the SEC three-year-comparative requirement, which are the provenance for the Rocket 2018 and 2019 rows.", "03_Case_Study_Data/Rocket Mortgage/Annual Reports/2020.pdf"),
        ("6", "Rocket Companies, Inc.", "10-K", "FY 2021 (filed 2022)", "Source for FY 2021 audited values used in Table A.2.", "03_Case_Study_Data/Rocket Mortgage/Annual Reports/2021.pdf"),
        ("7", "Rocket Companies, Inc.", "10-K", "FY 2022 (filed 2023)", "Source for FY 2022 audited values used in Table A.2.", "03_Case_Study_Data/Rocket Mortgage/Annual Reports/2022.pdf"),
        ("8", "Rocket Companies, Inc.", "10-K", "FY 2023 (filed 2024)", "Source for FY 2023 audited values used in Table A.2.", "03_Case_Study_Data/Rocket Mortgage/Annual Reports/2023.pdf"),
        ("9", "Rocket Companies, Inc.", "10-K", "FY 2024 (filed 2025)", "Source for FY 2024 audited values used in Table A.2; benchmark context throughout Chapter 5.", "03_Case_Study_Data/Rocket Mortgage/Annual Reports/2024.pdf"),
        ("10", "Rocket Companies, Inc.", "10-K", "FY 2025 (filed 2026)", "Source for FY 2025 audited values used in Table A.2; the FY 2025 disclosures reflect the Mr. Cooper acquisition, which is treated as a confounder for per-employee productivity comparison and discussed in §5.3.5.", "03_Case_Study_Data/Rocket Mortgage/Annual Reports/2025.pdf"),
    ]

    # Number of columns must match.
    if len(target.rows[0].cells) != len(new_rows[0]):
        # The current Table B.1 has 6 columns. Match.
        print(f"[B.1] adjusting column count from {len(target.rows[0].cells)} to {len(new_rows[0])}")

    while len(target.rows) > 1:
        target._tbl.remove(target.rows[1]._tr)

    # Update header row.
    for i, txt in enumerate(new_rows[0]):
        if i < len(target.rows[0].cells):
            cell = target.rows[0].cells[i]
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            cell.paragraphs[0].add_run(txt).bold = True

    for row in new_rows[1:]:
        new_row = target.add_row().cells
        for i, txt in enumerate(row):
            if i < len(new_row):
                cell = new_row[i]
                for p in cell.paragraphs:
                    for r in list(p.runs):
                        r._element.getparent().remove(r._element)
                cell.paragraphs[0].add_run(str(txt))

    print(f"[B.1] rewritten with {len(new_rows)-1} filings")
    return True


# ---------------------------------------------------------------------------
# Table D.1 — Top-10 BERTopic topics (real labels and shares)
# ---------------------------------------------------------------------------

def rewrite_table_d1(doc):
    df = pd.read_csv(CSV_TOPICS)
    clustered = df[df["topic"] != -1]
    n_clustered = len(clustered)
    pre = clustered[clustered["post_betsy"] == 0]
    post = clustered[clustered["post_betsy"] == 1]

    counts = (
        clustered.groupby(["topic", "topic_label"]).size().reset_index(name="n")
        .sort_values("n", ascending=False)
    )
    rows = []
    for _, r in counts.head(10).iterrows():
        tid = r["topic"]
        share = r["n"] / n_clustered * 100
        p_pre = (pre["topic"] == tid).sum() / max(1, len(pre)) * 100
        p_post = (post["topic"] == tid).sum() / max(1, len(post)) * 100
        rows.append((
            int(tid),
            r["topic_label"],
            f"{share:.2f}",
            f"{p_pre:.2f}",
            f"{p_post:.2f}",
            f"{p_post - p_pre:+.2f}",
        ))

    target = None
    for t in doc.tables:
        if not t.rows:
            continue
        h = [c.text.strip() for c in t.rows[0].cells]
        if h[:1] == ["Topic ID"]:
            target = t
            break
    if target is None:
        print("[D.1] table not found")
        return False

    new_header = [
        "Topic ID", "BERTopic label", "Clustered share (%)",
        "Pre-Betsy share (%)", "Post-Betsy share (%)", "Δpp (post − pre)",
    ]
    while len(target.rows) > 1:
        target._tbl.remove(target.rows[1]._tr)

    # Adapt column count: original has 5 cols. Need 6.
    while len(target.rows[0].cells) < len(new_header):
        # python-docx does not expose a clean add_column; we duplicate the
        # last cell in the header row to grow the table column count.
        from docx.oxml.ns import qn
        last_tc = target.rows[0].cells[-1]._tc
        new_tc = last_tc.makeelement(qn("w:tc"), {})
        # Copy properties.
        from copy import deepcopy
        new_tc = deepcopy(last_tc)
        last_tc.addnext(new_tc)

    for i, txt in enumerate(new_header):
        if i < len(target.rows[0].cells):
            cell = target.rows[0].cells[i]
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            cell.paragraphs[0].add_run(txt).bold = True

    for row in rows:
        new_row = target.add_row().cells
        for i, txt in enumerate(row):
            if i < len(new_row):
                cell = new_row[i]
                for p in cell.paragraphs:
                    for r in list(p.runs):
                        r._element.getparent().remove(r._element)
                cell.paragraphs[0].add_run(str(txt))

    print(f"[D.1] rewritten with top-10 BERTopic topics (n_clustered = {n_clustered})")
    return True


# ---------------------------------------------------------------------------
# Table D.2 — Keyword prevalence (the actual seven groups)
# ---------------------------------------------------------------------------

def rewrite_table_d2(doc):
    import scipy.stats as ss
    df = pd.read_csv(CSV_KW)
    pre = df[df["post_betsy"] == 0]
    post = df[df["post_betsy"] == 1]
    n_pre = len(pre)
    n_post = len(post)

    groups = [
        ("Explicit AI / automation", "ai, bot, chatbot, automated, betsy", "Explicit_AI_Automation"),
        ("Speed / efficiency", "fast, quick, slow, delay, wait", "Speed_Efficiency"),
        ("Human support", "agent, officer, human, broker, person", "Human_Support"),
        ("Communication", "respond, communication, email, call", "Communication"),
        ("Delay / friction", "delay, hold, wait, stuck, friction", "Delay_Friction"),
        ("Cost / rates", "rate, fee, cost, expensive, pricing", "Cost_Rates"),
        ("Digital / process", "online, app, portal, dashboard, login", "Digital_Process"),
    ]

    bonferroni_alpha = 0.05 / len(groups)
    rows = []
    for label, keywords, col in groups:
        p_pre = pre[col].mean() * 100
        p_post = post[col].mean() * 100
        delta = p_post - p_pre
        pre_count = pre[col].sum()
        post_count = post[col].sum()
        table = [[pre_count, n_pre - pre_count], [post_count, n_post - post_count]]
        chi2, p, _, _ = ss.chi2_contingency(table, correction=True)
        rows.append((
            label,
            keywords,
            f"{p_pre:.2f}",
            f"{p_post:.2f}",
            f"{delta:+.2f}",
            f"{p:.4f}" if p >= 0.0001 else f"{p:.1e}",
            "Yes" if p < bonferroni_alpha else "No",
        ))

    target = None
    for t in doc.tables:
        if not t.rows:
            continue
        h = [c.text.strip() for c in t.rows[0].cells]
        if h[:1] == ["Group"]:
            target = t
            break
    if target is None:
        print("[D.2] table not found")
        return False

    new_header = [
        "Keyword group", "Representative keywords",
        "Pre prevalence (%)", "Post prevalence (%)", "Δpp",
        "p-value (χ², Yates)", f"Bonferroni-significant (α/7 = {bonferroni_alpha:.5f})",
    ]
    while len(target.rows) > 1:
        target._tbl.remove(target.rows[1]._tr)

    for i, txt in enumerate(new_header):
        if i < len(target.rows[0].cells):
            cell = target.rows[0].cells[i]
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            cell.paragraphs[0].add_run(txt).bold = True

    for row in rows:
        new_row = target.add_row().cells
        for i, txt in enumerate(row):
            if i < len(new_row):
                cell = new_row[i]
                for p in cell.paragraphs:
                    for r in list(p.runs):
                        r._element.getparent().remove(r._element)
                cell.paragraphs[0].add_run(str(txt))

    print(f"[D.2] rewritten with {len(rows)} keyword groups (Bonferroni α/{len(groups)} = {bonferroni_alpha:.5f})")
    return True


# ---------------------------------------------------------------------------
# Body-text reconciliation in Appendix A.1, Appendix C.1, Appendix C.2,
# Appendix D captions, and methodology cross-references
# ---------------------------------------------------------------------------

def fix_appendix_a1_overview(doc):
    """Rewrite the Appendix A.1 overview paragraph to match the CSV row count."""
    target = None
    for p in doc.paragraphs:
        if "thirteen annual observations" in p.text or "13 annual observations" in p.text or "five for Better" in p.text:
            target = p
            break
    if target is None:
        for i, p in enumerate(doc.paragraphs):
            if "A.1 Dataset Overview" in p.text and i + 1 < len(doc.paragraphs):
                target = doc.paragraphs[i + 1]
                break
    if target is None:
        print("[A.1] overview paragraph not found")
        return False

    new_text = (
        "The financial dataset analysed in Chapter 5 contains thirteen audited "
        "annual observations: five for Better Home & Finance Holding Company "
        "(formerly Better HoldCo, Inc.) covering fiscal years 2021 through 2025, "
        "and eight for Rocket Companies, Inc. covering fiscal years 2018 through "
        "2025. The Better 2021 row is sourced from the S-1 historical financials "
        "because the 2023 10-K, which Better files as a Smaller Reporting Company, "
        "discloses only two years of audited income-statement data. The Rocket "
        "2018 and 2019 rows are sourced from the audited three-year comparative "
        "tables presented in the FY 2020 10-K (Rocket Companies, 2021), which is "
        "the earliest filing available on disk. All revenue, expense, and net "
        "income / (loss) values are reported in millions of US dollars."
    )
    rewrite_paragraph_text_preserve_format(target, new_text)
    print("[A.1] dataset-overview paragraph rewritten")
    return True


def fix_appendix_c_extraction(doc):
    """Replace the Appendix C extraction-date paragraph to match the actual
    scrape timestamp in Better_Reviews.csv (April 26, 2026)."""
    target = None
    for p in doc.paragraphs:
        if "May 6, 2026" in p.text or ("Trustpilot" in p.text and "Apify" in p.text and "extracted" in p.text):
            target = p
            break
    if target is None:
        print("[C.1] extraction paragraph not found")
        return False

    new_text = (
        "The Trustpilot review corpus was extracted on a single date "
        "(26 April 2026) using the Apify Trustpilot Review Scraper "
        "(Actor ID l3wcDhSSC96LBRUpc), targeting the better.com domain. "
        "The raw extraction returned 1,934 reviews spanning 22 August 2020 "
        "to 25 April 2026 and is archived as a static dataset in "
        "Data/NLP/Better_Reviews.csv. The thesis title-page date (April 2026) "
        "and the Trustpilot extraction date are therefore in the same calendar "
        "month, and the empirical cut-off is set at the latest review date "
        "(25 April 2026)."
    )
    rewrite_paragraph_text_preserve_format(target, new_text)
    print("[C.1] extraction-date paragraph rewritten (now: 26 April 2026)")
    return True


def fix_appendix_c_cleaning(doc):
    """Replace the Appendix C cleaning paragraph with the sequence actually
    executed in cleaning_reviews.ipynb. The verified counts are:

        raw 1,934 -> dropna(0) -> drop_duplicates(0) -> length>20(-18) -> 'en'(-1) -> 1,915

    """
    target = None
    in_c2 = False
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 2" and ("C.2" in p.text or "Cleaning Pipeline" in p.text):
            in_c2 = True
            continue
        if in_c2 and p.style.name in ("Heading 1", "Heading 2"):
            break
        if in_c2 and ("Cleaning steps" in p.text or "exact-duplicate" in p.text or "removal of" in p.text):
            target = p
            break
    if target is None:
        # Fallback: locate by content.
        for p in doc.paragraphs:
            if "exact-duplicate" in p.text and "Better-controlled" in p.text:
                target = p
                break
    if target is None:
        print("[C.2] cleaning paragraph not found")
        return False

    new_text = (
        "The cleaning pipeline is documented in 08_Code/cleaning_reviews.ipynb "
        "and executes the following sequential filters on the raw extraction of "
        "1,934 reviews (counts verified by re-executing the notebook against "
        "Data/NLP/Better_Reviews.csv): "
        "(i) drop rows with missing review_text, rating, or review_date "
        "(0 dropped); "
        "(ii) drop duplicate review_id values (0 dropped, because Trustpilot "
        "review IDs are globally unique); "
        "(iii) drop reviews with body length of 20 characters or fewer (18 "
        "dropped); "
        "(iv) retain only reviews with the platform-tagged language metadata "
        "field equal to 'en' (1 non-English review dropped). "
        "The cleaned corpus contains 1,915 reviews. The pre/post split is "
        "anchored to the Betsy launch date (17 October 2024): 1,746 reviews "
        "pre-Betsy, 169 reviews post-Betsy. The wording in §3.4.2 of Chapter 3 "
        "describes the same four-step pipeline and reports the same final "
        "corpus counts."
    )
    rewrite_paragraph_text_preserve_format(target, new_text)
    print("[C.2] cleaning paragraph rewritten (now: 4 documented steps)")
    return True


def fix_appendix_d2_caption(doc):
    """Update the D.2 caption: Bonferroni threshold should be 0.05/7 = 0.00714,
    not 0.05/11 = 0.00455 (the analysis uses seven keyword groups, not eleven).
    """
    target = None
    for p in doc.paragraphs:
        if p.style.name == "Caption" and "Keyword-group prevalence" in p.text:
            target = p
            break
    if target is None:
        print("[D.2 caption] not found")
        return False
    new_text = (
        "Table D.2. Keyword-group prevalence in the cleaned 1,915-review "
        "Trustpilot corpus, pre vs post Betsy. Prevalence is the share of "
        "reviews containing at least one keyword from the group. p-values are "
        "from two-sided chi-square tests with the Yates correction. The "
        "Bonferroni-corrected significance threshold for seven keyword groups "
        "is α/7 = 0.05/7 ≈ 0.00714. Cost / rates is the only group whose "
        "pre-to-post change survives this correction."
    )
    rewrite_paragraph_text_preserve_format(target, new_text)
    print("[D.2 caption] updated")
    return True


def fix_appendix_d1_caption(doc):
    target = None
    for p in doc.paragraphs:
        if p.style.name == "Caption" and "BERTopic topics" in p.text:
            target = p
            break
    if target is None:
        print("[D.1 caption] not found")
        return False
    new_text = (
        "Table D.1. Top ten BERTopic topics by clustered-corpus share. Topic "
        "identifiers and labels are produced by the BERTopic pipeline "
        "documented in §C.4. Shares are computed within the clustered subset "
        "(n_clustered = 1,198 of 1,804 reviews used as topic-modelling input; "
        "the residual 606 are HDBSCAN noise points). Pre / post shares are "
        "computed within the pre-Betsy clustered subset (n = 1,077) and the "
        "post-Betsy clustered subset (n = 121). Δpp is the percentage-point "
        "difference (post − pre)."
    )
    rewrite_paragraph_text_preserve_format(target, new_text)
    print("[D.1 caption] updated")
    return True


def rewrite_paragraph_text_preserve_format(paragraph, new_text: str) -> None:
    if paragraph.runs:
        r0 = paragraph.runs[0]
        font_name = r0.font.name
        font_size = r0.font.size
        bold = r0.bold
        italic = r0.italic
    else:
        font_name = font_size = bold = italic = None
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    nr = paragraph.add_run(new_text)
    if font_name:
        nr.font.name = font_name
    if font_size:
        nr.font.size = font_size
    if bold is not None:
        nr.bold = bold
    if italic is not None:
        nr.italic = italic


def main():
    doc = open_draft()
    rewrite_table_a1(doc)
    rewrite_table_a2(doc)
    rewrite_table_b1(doc)
    rewrite_table_d1(doc)
    rewrite_table_d2(doc)
    fix_appendix_a1_overview(doc)
    fix_appendix_c_extraction(doc)
    fix_appendix_c_cleaning(doc)
    fix_appendix_d1_caption(doc)
    fix_appendix_d2_caption(doc)
    save_draft(doc)
    print("Saved.")


if __name__ == "__main__":
    main()
