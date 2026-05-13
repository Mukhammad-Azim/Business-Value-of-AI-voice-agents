"""
insert_chapter5.py
==================

Insert Chapter 5 ("Empirical Findings: Triangulated Evidence on the Business
Value of Voice AI at Better.com") into ``04_Drafts/Main_Draft.docx`` between
the existing Chapter 4 (Case Study Description) and the References list.

The chapter follows the building plan in
``Chapter_5_Building_Plan_Master_Thesis.docx`` (provided by the supervisor)
and uses only the standardized financial / NLP values fixed during the
pre-Chapter-5 cleanup pass:

* Better revenue: $76.8M (2023) -> $108.5M (2024) -> $164.9M (2025).
* Better total expenses (standardized base): $368.1M / $313.9M / $330.7M.
* Better net loss: -$536.4M / -$206.3M / -$165.9M.
* Expense ratio: 4.79 / 2.89 / 2.01.
* Revenue per employee: $0.094M / $0.087M / $0.124M
  (pre-period 2021-2023 average $0.130M).
* Operational Leverage Index: 195.3 in 2025, 2023 = 100 (appendix only).
* NLP corpus: 1,934 raw -> 1,915 clean (Pre 1,746 / Post 169) ->
  1,804 topic-modeling input -> 1,198 clustered + 606 noise; clustered
  non-outlier denominators Pre 1,077 / Post 121.
* Star ratings: 4.2806 -> 3.8580 (Welch p = 0.0019, Cohen's d = 0.2899).
* VADER sentiment: 0.6416 -> 0.5151 (p = 0.0057, Cohen's d = 0.2505).
* Keyword prevalences exactly as in the supervisor brief.

The chapter is written in the cautious "consistent with" / "associated with"
register the supervisor brief mandates. No claim of causal identification.
"""

from __future__ import annotations

import os
import sys
from copy import deepcopy

import docx
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DRAFT_PATH = os.path.join(REPO_ROOT, "04_Drafts", "Main_Draft.docx")
FIG_FIN = os.path.join(REPO_ROOT, "outputs", "figures", "financial")
FIG_NLP = os.path.join(REPO_ROOT, "outputs", "figures", "NLP")


# ---------------------------------------------------------------------------
# Helpers for inserting paragraphs at a specific anchor (instead of appending
# to the end of the document, which is python-docx's default behavior).
# ---------------------------------------------------------------------------


def insert_paragraph_before(anchor_para, text: str = "", style_name: str | None = None):
    """Insert a new paragraph (with given text + style) directly before the
    anchor paragraph and return the new Paragraph object.

    Implementation: append a paragraph to the document body so that python-
    docx attaches a proper parent, then move the underlying ``w:p`` element
    to the position immediately before the anchor.
    """
    doc = anchor_para.part.document
    tmp = doc.add_paragraph()
    if style_name:
        tmp.style = doc.styles[style_name]
    if text:
        tmp.add_run(text)
    # Move the new paragraph's XML to just before the anchor.
    anchor_para._p.addprevious(tmp._p)
    return tmp


def insert_picture_before(anchor_para, image_path: str, width_in: float = 6.0):
    """Insert a paragraph containing an image, directly before the anchor
    paragraph. Returns the new Paragraph."""
    doc = anchor_para.part.document
    tmp = doc.add_paragraph()
    run = tmp.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    # Center-align figure
    tmp.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    anchor_para._p.addprevious(tmp._p)
    return tmp


# ---------------------------------------------------------------------------
# Chapter 5 content
# ---------------------------------------------------------------------------


# Each entry is (kind, payload):
#   ("h1", text)   -> Heading 1
#   ("h2", text)   -> Heading 2
#   ("h3", text)   -> Heading 3
#   ("p",  text)   -> Normal paragraph
#   ("fig", (relpath, caption_text)) -> figure + Caption paragraph
#   ("quote", text) -> Quote-styled paragraph

CONTENT: list[tuple[str, object]] = []


def H1(t):
    CONTENT.append(("h1", t))


def H2(t):
    CONTENT.append(("h2", t))


def H3(t):
    CONTENT.append(("h3", t))


def P(t):
    CONTENT.append(("p", t))


def FIG(path, caption):
    CONTENT.append(("fig", (path, caption)))


# ---------------------------------------------------------------------------
# 5.0 Chapter heading
# ---------------------------------------------------------------------------

H1(
    "5. Empirical Findings: Triangulated Evidence on the Business Value of "
    "Voice AI at Better.com"
)

# ---------------------------------------------------------------------------
# 5.1 Introduction to Empirical Findings
# ---------------------------------------------------------------------------

H2("5.1 Introduction to Empirical Findings")

P(
    "This chapter presents the empirical findings of the thesis by integrating "
    "three evidence streams: qualitative case-study evidence on the integration "
    "of Betsy into Better.com\u2019s Tinman-based mortgage origination workflow, "
    "longitudinal financial evidence drawn from SEC filings and a descriptive "
    "benchmark comparison against Rocket Mortgage, and customer-perception "
    "evidence derived from a Natural Language Processing (NLP) analysis of "
    "Trustpilot reviews. The chapter is organized around the two research "
    "questions of the thesis. RQ1 asks how voice LLM agents have been integrated "
    "into Better.com\u2019s mortgage origination process and what architectural, "
    "workflow, and human-AI collaboration prerequisites this integration "
    "implies. RQ2 asks what measurable business value can be associated with "
    "Better.com\u2019s voice AI deployment across operational efficiency, "
    "productivity, customer experience, and financial performance, and what "
    "factors mediate the realization of this value."
)

P(
    "RQ1 is answered primarily through the qualitative case-study evidence "
    "established in Chapter 4, supplemented by company disclosures and "
    "vendor materials (Better Home & Finance Holding Company, 2024a, 2026; "
    "ElevenLabs, 2026). RQ2 is answered through two largely independent "
    "quantitative streams: the longitudinal financial analysis of Better\u2019s "
    "2021\u20132025 SEC filings, contextualized against Rocket Mortgage as a "
    "directional industry benchmark; and the NLP analysis of 1,915 cleaned "
    "Trustpilot reviews segmented around Betsy\u2019s October 17, 2024 launch "
    "(Better Home & Finance Holding Company, 2024a). Following the convergent "
    "parallel mixed-methods design described in Chapter 3, the three streams "
    "are interpreted together, with explicit attention to where they "
    "converge and where they diverge."
)

P(
    "Two interpretive constraints frame the chapter. First, the thesis relies "
    "on observational, public, secondary data; consequently, the chapter does "
    "not claim causal identification of Betsy\u2019s independent contribution. "
    "It evaluates whether the post-deployment evidence is consistent with "
    "AI-enabled operational leverage in the sense theorized by Brynjolfsson "
    "et al. (2023) and Aggarwal et al. (2025), and how customer-discourse "
    "evidence supports, complicates, or qualifies that interpretation. Second, "
    "the chapter foregrounds an analytically valuable divergence: financial "
    "and operational metrics improve materially across 2024\u20132025, while "
    "independent Trustpilot star ratings and VADER sentiment decline "
    "significantly over the same period and explicit AI mentions remain "
    "near-zero. Rather than treating this divergence as a contradiction, the "
    "chapter treats it as a triangulation insight: different measurement "
    "instruments capture different facets of customer experience, and the "
    "business value of voice AI in financial services appears to be "
    "multidimensional and mediated rather than uniformly directional."
)

P(
    "The chapter proceeds in six analytical sections. Section 5.2 addresses "
    "RQ1 by analyzing how Tinman, Betsy, ElevenLabs, and human escalation "
    "fit together as an integration mechanism. Sections 5.3 and 5.4 address "
    "the operational/financial and customer-experience dimensions of RQ2. "
    "Section 5.5 synthesizes the three streams into a triangulated "
    "interpretation. Section 5.6 records the boundary conditions of the "
    "empirical findings, and Section 5.7 presents direct, evidence-bounded "
    "answers to RQ1 and RQ2 and transitions to the design-pattern argument "
    "developed in Chapter 6."
)

# ---------------------------------------------------------------------------
# 5.2 RQ1 - Integration of Voice LLM Agents
# ---------------------------------------------------------------------------

H2("5.2 RQ1: Integration of Voice LLM Agents into Mortgage Origination")

P(
    "RQ1 asks how voice LLM agents have been integrated into Better.com\u2019s "
    "mortgage origination process and under what prerequisites this "
    "integration becomes operationally meaningful. The case evidence developed "
    "in Chapter 4 supports four interlocking observations: Tinman is the "
    "architectural prerequisite that makes a conversational layer "
    "operationally useful; Betsy is best understood as a coordination layer "
    "rather than a stand-alone underwriter; ElevenLabs supplies the voice "
    "infrastructure but does not own Better\u2019s mortgage logic; and the "
    "deployment is structured as task reallocation between AI and licensed "
    "human specialists, not as full labor substitution. Together, these "
    "observations indicate that voice LLM integration in financial services "
    "is best interpreted as a workflow-integration problem rather than a "
    "front-end chatbot problem."
)

H3("5.2.1 Tinman as the Architectural Prerequisite")

P(
    "Better.com\u2019s ability to deploy a voice AI loan assistant rests on the "
    "fact that its mortgage process is already encoded in Tinman, the "
    "company\u2019s proprietary, integrated mortgage operating platform. Tinman "
    "functions as the system of record for borrower data, pricing, "
    "documentation, eligibility checks, workflow routing, and audit trails "
    "(Better Home & Finance Holding Company, 2026). This matters because, "
    "in the task-technology fit tradition of Goodhue and Thompson (1995), "
    "the fit between a conversational technology and the work it is asked to "
    "perform depends on whether the technology has access to the data and "
    "process state needed for the task. A voice LLM agent operating without "
    "access to a structured mortgage system would be limited to generic "
    "informational dialogue. Operating on top of Tinman, by contrast, the "
    "agent can plausibly reference an applicant\u2019s loan stage, request a "
    "specific missing document, or surface a status update derived from the "
    "system of record. The case evidence therefore suggests that the "
    "operational meaningfulness of Betsy is not a property of the voice "
    "agent in isolation; it is a joint property of Betsy and the "
    "Tinman-mediated workflow into which it is embedded."
)

P(
    "This interpretation is consistent with the broader literature on AI in "
    "customer-facing financial services. Hentzen et al. (2022) emphasize that "
    "AI systems in mortgage and banking contexts are most effective when "
    "embedded in compliant, auditable workflows; Mogaji and Nguyen (2021) "
    "note that consequential AI applications often operate as backend or "
    "infrastructure-level components rather than as branded customer-facing "
    "chatbots. Better.com\u2019s public materials position Tinman as such an "
    "architectural foundation: the company describes itself as a "
    "technology-first mortgage lender whose digital infrastructure was "
    "designed to digitize the end-to-end home-finance process before voice "
    "AI was layered on top (Better.com, n.d.; Better Home & Finance Holding "
    "Company, 2026). The case therefore supports the argument that the "
    "preconditions for a meaningful voice LLM deployment include not only "
    "model quality but also a process-aware operating platform that "
    "exposes structured workflow state to the conversational layer."
)

H3("5.2.2 Betsy as a Conversational Coordination Layer")

P(
    "Betsy is described in Better.com\u2019s investor and press materials as a "
    "voice-based AI loan assistant rather than as an underwriting engine "
    "(Better Home & Finance Holding Company, 2024a; ElevenLabs, 2026). The "
    "publicly disclosed scope of Betsy\u2019s tasks centers on borrower "
    "communication: handling inbound and outbound calls, answering routine "
    "questions about rates and processes, collecting and chasing missing "
    "documentation, providing application-status updates, and triaging "
    "borrower inquiries. In this sense, Betsy is best interpreted as a "
    "conversational coordination layer over the mortgage workflow. It "
    "automates the high-volume, repetitive communicative work that "
    "historically consumed loan-officer time, while leaving credit "
    "decisions, complex problem solving, and final closing steps to "
    "licensed human specialists."
)

P(
    "This division of labor is consistent with Huang and Rust\u2019s (2018) "
    "service-AI hierarchy, which distinguishes mechanical, analytical, "
    "intuitive, and empathetic dimensions of service work. Betsy operates "
    "primarily within the mechanical and analytical dimensions \u2014 retrieving "
    "structured information, executing predictable conversational scripts, "
    "and handling routine status communication \u2014 while higher-stakes "
    "interactions involving borrower-specific judgment, regulatory advice, "
    "or escalation are routed to humans. It is also broadly consistent with "
    "the field-experimental evidence from Wang et al. (2023) on voice-based "
    "AI in call-center settings, where voice AI agents complement rather "
    "than replace human workers, and from Zhang et al. (2023) on the "
    "operational performance effects of AI-based conversational agents "
    "in call centers. Crucially, however, the case-study evidence does "
    "not allow attribution of any specific operational outcome at "
    "Better.com to Betsy in isolation; it shows that Betsy is positioned "
    "as a coordination layer within a broader Betsy/Tinman operating model."
)

H3("5.2.3 ElevenLabs as Voice Infrastructure")

P(
    "ElevenLabs functions as the underlying voice-agent infrastructure for "
    "Betsy. According to the joint Better.com\u2013ElevenLabs materials, "
    "ElevenLabs provides the conversational AI agents (\u201cElevenAgents\u201d), "
    "low-latency speech synthesis, and the agentic conversational runtime "
    "that lets Betsy speak naturally with borrowers (ElevenLabs, 2026; "
    "Better Home & Finance Holding Company & ElevenLabs, 2026). This "
    "vendor relationship is analytically important for two reasons. First, "
    "it clarifies that the voice layer is delivered through a third-party "
    "platform rather than constructed in-house, which aligns with broader "
    "industry patterns in which financial-services firms compose "
    "AI-supported workflows from specialized infrastructure components "
    "(Liu et al., 2024). Second, it bounds the analytical claims that can "
    "be made: ElevenLabs supplies the voice-agent capabilities, but it "
    "does not own the mortgage logic, the customer relationship, or the "
    "regulated decision-making. Better.com retains those through Tinman "
    "and through its licensed loan officers."
)

P(
    "The vendor evidence must therefore be treated cautiously. ElevenLabs "
    "and Better.com both have incentives to present the deployment "
    "favorably in joint case-study materials, and figures cited in vendor "
    "communications are not independently audited. Following Mogaji and "
    "Nguyen\u2019s (2021) caution about treating AI vendor claims as "
    "independent evidence of business value, this thesis uses the "
    "Better.com\u2013ElevenLabs materials to characterize the integration "
    "mechanism rather than to substantiate magnitude claims about cost "
    "savings or revenue uplift. Where vendor materials and Better\u2019s own "
    "SEC disclosures align (for example, on Betsy\u2019s role as a voice-based "
    "loan assistant), the alignment is reported; where vendor materials "
    "extend beyond what Better\u2019s SEC filings disclose, the claim is "
    "treated as moderate-strength descriptive evidence only."
)

H3("5.2.4 Human-AI Collaboration and Regulated Task Boundaries")

P(
    "A consistent feature of the public materials is that Betsy\u2019s "
    "operating model preserves human escalation. Better Home & Finance "
    "Holding Company (2024a) describes Betsy as an assistant that handles "
    "routine borrower communication and surfaces escalations to licensed "
    "loan officers, and the joint Better\u2013ElevenLabs materials emphasize "
    "compliance-oriented controls around the conversational layer "
    "(ElevenLabs, 2026; Better Home & Finance Holding Company & ElevenLabs, "
    "2026). The case evidence therefore supports an interpretation of the "
    "deployment as task reallocation rather than full labor substitution: "
    "high-volume informational and procedural communication is reallocated "
    "to a voice AI agent, while consequential, regulated, or borrower-"
    "specific judgments remain with human specialists. This pattern is "
    "consistent with Acemoglu and Restrepo\u2019s (2019) framework, in which "
    "automation simultaneously displaces labor on some tasks and "
    "reinstates labor on adjacent tasks that the technology cannot perform "
    "as effectively."
)

P(
    "From a regulatory perspective, mortgage origination in the United "
    "States is governed by federal and state-level disclosure, fair-lending, "
    "and consumer-protection rules administered through agencies such as "
    "the Consumer Financial Protection Bureau (Consumer Financial Protection "
    "Bureau, 2024, 2025). Many of these rules ultimately require licensed "
    "personnel to validate or sign off on consequential decisions and "
    "disclosures. A voice LLM agent integrated into a regulated origination "
    "workflow must therefore be embedded in an architecture that maintains "
    "auditable handoffs, documented decision trails, and visible "
    "human-in-the-loop checkpoints. The Better.com case is consistent with "
    "this configuration: Tinman maintains the workflow and audit trails, "
    "Betsy mediates routine communication, and licensed loan officers "
    "retain authority over consequential decisions."
)

H3("5.2.5 Integration Prerequisites: Synthesis for RQ1")

P(
    "Synthesizing the case evidence, RQ1 can be answered as follows. Voice "
    "LLM agents become operationally meaningful in mortgage origination "
    "when they are integrated into a process-aware operating platform that "
    "exposes structured workflow state, when they are scoped to a "
    "communication-coordination role rather than a regulated decision-making "
    "role, when they sit on top of low-latency voice infrastructure capable "
    "of natural conversation, and when they are governed by an explicit "
    "human-AI collaboration model with auditable escalation paths. In the "
    "Better.com case, Tinman supplies the first prerequisite, Betsy "
    "operationalizes the second, ElevenLabs supplies the third, and the "
    "regulated task structure of mortgage origination supplies the fourth. "
    "These prerequisites are case-derived and therefore most directly "
    "applicable to firms whose business processes can be analogously "
    "encoded in a digital operating platform. They are nonetheless "
    "consistent with the broader literature on task-technology fit "
    "(Goodhue & Thompson, 1995), AI-augmented service work (Huang & Rust, "
    "2018), and AI in customer-facing financial services (Hentzen et al., "
    "2022; Mogaji & Nguyen, 2021)."
)

# ---------------------------------------------------------------------------
# 5.3 RQ2 - Operational and Financial Business Value
# ---------------------------------------------------------------------------

H2("5.3 RQ2: Operational and Financial Business Value")

P(
    "This section addresses the operational and financial dimensions of RQ2. "
    "The analytical strategy is descriptive and longitudinal. The 2024-to-2025 "
    "comparison is treated as the primary post-transition window, because "
    "Betsy launched on October 17, 2024 and 2025 is therefore the first full "
    "post-launch fiscal year (Better Home & Finance Holding Company, 2024a). "
    "The 2023-to-2025 comparison is reported as supporting evidence of "
    "recovery from the 2023 trough rather than as a Betsy-specific causal "
    "comparison. The pre-period 2021\u20132023 average is used as a robustness "
    "check, to show that 2025 revenue and revenue per employee, while "
    "materially higher than the 2024 trough, have not fully returned to "
    "pandemic-era levels. Throughout, Rocket Companies is used as a "
    "directional industry benchmark; differences in scale, business model, "
    "and the 2025 Mr. Cooper acquisition mean that Rocket cannot serve as a "
    "formal counterfactual."
)

# Figure 5.1 - Core financial trajectory
FIG(
    os.path.join(FIG_FIN, "fig_fin_better_financial_trajectory_clean.png"),
    "Figure 5.1. Better.com\u2019s revenue, standardized total expense base, "
    "and net income/loss, 2021\u20132025. The shaded periods distinguish the "
    "pre-Betsy baseline (2021\u20132023), the 2024 transition year, and the 2025 "
    "post-Betsy fiscal year. Source: Author\u2019s construction from Better Home "
    "& Finance Holding Company SEC filings (S-1; 10-K filings for fiscal "
    "years 2023, 2024, and 2025). Figure provides descriptive evidence of "
    "revenue recovery, expense containment, and narrowing losses; it does "
    "not isolate Betsy\u2019s causal contribution.",
)

H3("5.3.1 Revenue Recovery and Post-Transition Growth")

P(
    "Better.com\u2019s revenue trajectory across 2023\u20132025 displays a clear "
    "recovery pattern. Reported total net revenue rose from approximately "
    "$76.8 million in 2023 to approximately $108.5 million in 2024 and "
    "approximately $164.9 million in 2025 (Better Home & Finance Holding "
    "Company, 2024b, 2025a, 2026a). The headline post-transition comparison "
    "is therefore an increase of approximately 52.0 percent from 2024 to "
    "2025. Across the longer 2023-to-2025 window, revenue grew by "
    "approximately 114.6 percent, but this larger figure incorporates the "
    "trough-year baseline of 2023, when the U.S. mortgage market was "
    "constrained by elevated interest rates and depressed origination "
    "volumes (Mortgage Bankers Association, 2024). Funded loan volume "
    "moved in the same direction, rising from approximately $3.0 billion "
    "in 2023 to approximately $3.6 billion in 2024 (Better Home & Finance "
    "Holding Company, 2024b, 2025a). Figure 5.1 visualizes this trajectory "
    "alongside the standardized total expense base and the net loss line."
)

P(
    "The revenue trajectory is consistent with the operational-leverage "
    "interpretation that Better\u2019s broader Betsy/Tinman AI operating model "
    "is associated with growth that does not require proportional expansion "
    "of the cost base. It is not, however, evidence that voice AI alone "
    "produced the recovery. The same period encompasses concurrent product "
    "developments (notably the launch of the NEO channel in early 2025), "
    "ongoing restructuring, the maturation of Tinman, and a partial "
    "recovery in the broader U.S. mortgage market that affected most "
    "originators (Mortgage Bankers Association, 2024). The most defensible "
    "framing is therefore that the post-transition revenue recovery is "
    "consistent with AI-enabled operational leverage, while leaving the "
    "specific contribution of Betsy unresolved."
)

H3("5.3.2 Expense Intensity and Operational Efficiency")

P(
    "The strongest descriptive evidence for operational efficiency at "
    "Better.com is the trajectory of the expense ratio, defined here as "
    "the standardized total expense base divided by total net revenue. "
    "On this construction, the expense ratio improved from 4.79 in 2023 to "
    "2.89 in 2024 and 2.01 in 2025. Equivalently, between 2024 and 2025 "
    "revenue grew by approximately 52.0 percent while the standardized "
    "expense base grew by only approximately 5.3 percent (from "
    "approximately $313.9 million in 2024 to approximately $330.7 million "
    "in 2025). This pattern \u2014 a substantial step up in revenue with only "
    "marginal growth in expenses \u2014 is the textbook descriptive signature "
    "of operational leverage and is consistent with the AI-enabled "
    "operational-leverage hypothesis examined in this thesis."
)

# Figure 5.2 - Expense ratio
FIG(
    os.path.join(FIG_FIN, "fig_fin_expense_ratio_cashflow_clean.png"),
    "Figure 5.2. Better.com\u2019s expense ratio (standardized total expense "
    "base / total net revenue) and operating cash flow, 2021\u20132025. The "
    "expense-ratio trajectory is the principal descriptive indicator of "
    "expense intensity; the operating cash flow panel is included to show "
    "that 2025 expense-ratio improvements are not yet matched by sustained "
    "positive operating cash flow. Source: Author\u2019s construction from "
    "Better Home & Finance Holding Company SEC filings (S-1; 10-K filings "
    "for fiscal years 2023, 2024, and 2025).",
)

P(
    "Two interpretive caveats apply. First, the construction of the "
    "standardized total expense base is consistent within Better\u2019s "
    "reported financial statements, but Better and Rocket use different "
    "accounting presentations (Rocket Companies, 2024, 2025). Cross-company "
    "expense-ratio comparisons should therefore be interpreted directionally, "
    "not as precise like-for-like measurements. Second, the operating cash "
    "flow panel in Figure 5.2 makes clear that improvements in the expense "
    "ratio are not yet matched by sustained positive operating cash flow; "
    "Better remained operating-cash-flow negative through 2025. The "
    "expense-ratio improvement is therefore best read as evidence of "
    "intensifying operational leverage rather than as evidence that the "
    "cost-base improvements have already translated into self-funding "
    "operations."
)

H3("5.3.3 Net Loss Reduction and Profitability Trajectory")

P(
    "Better.com\u2019s net loss narrowed from approximately $536.4 million in "
    "2023 to approximately $206.3 million in 2024 and approximately $165.9 "
    "million in 2025 (Better Home & Finance Holding Company, 2024b, 2025a, "
    "2026a). On a year-over-year basis, the 2024-to-2025 net-loss "
    "improvement is approximately 19.6 percent. The direction is favorable "
    "and the magnitude of the multi-year improvement is substantial. At "
    "the same time, the company has not achieved profitability. Reading "
    "the 2025 result as a financial turnaround would therefore overstate "
    "the evidence: losses narrowed materially, but the company remained "
    "unprofitable in 2025, and the profit margin remained negative."
)

P(
    "The trajectory is consistent with the operational-leverage "
    "interpretation that revenue is growing faster than the cost base, "
    "while also suggesting that the cost base remains structurally large "
    "relative to revenue. The comparison with Rocket Companies (Section "
    "5.3.5) helps contextualize this trajectory: across 2023\u20132025, both "
    "originators experienced cyclical pressures and partial recovery, but "
    "Better\u2019s relative loss-narrowing is descriptively more pronounced. "
    "Because Better and Rocket differ substantially in scale, business "
    "mix, and 2025 acquisition context, this comparative observation is "
    "treated as directional rather than as a controlled estimate."
)

H3("5.3.4 Labor Productivity and Scalability")

P(
    "Revenue per employee (RPE), defined as total net revenue divided by "
    "year-end headcount, is used as the primary proxy for labor "
    "productivity and scalability. RPE moved from approximately $0.094 "
    "million in 2023 to approximately $0.087 million in 2024 and "
    "approximately $0.124 million in 2025. The 2024-to-2025 change is "
    "approximately +42.9 percent, indicating a substantial step up in "
    "revenue per worker during the first full post-Betsy fiscal year. "
    "Across the 2023-to-2025 window, the RPE improvement is approximately "
    "+32.3 percent. Headcount itself contracted sharply from the "
    "pandemic-era peak (approximately 10,400 in 2021) to a 2023 trough "
    "near 820, before rising modestly to approximately 1,250 in 2024 and "
    "1,329 in 2025 (Better Home & Finance Holding Company, 2024b, 2025a, "
    "2026a). The post-transition RPE improvement is therefore not driven "
    "primarily by further headcount contraction but by faster revenue "
    "growth on an essentially flat headcount base."
)

# Figure 5.3 - Headcount and RPE
FIG(
    os.path.join(FIG_FIN, "fig_fin_headcount_productivity_clean.png"),
    "Figure 5.3. Better.com\u2019s year-end headcount and revenue per employee "
    "(RPE), 2021\u20132025. The figure shows the post-2021 workforce contraction "
    "and the partial recovery of RPE in 2025. Source: Author\u2019s construction "
    "from Better Home & Finance Holding Company SEC filings. The figure "
    "supports the productivity dimension of business value but also shows "
    "that 2025 RPE remains below the 2021\u20132023 pre-period average of "
    "approximately $0.130 million.",
)

P(
    "The productivity evidence is moderately strong but should be qualified. "
    "Although 2025 RPE is materially higher than 2024 RPE, it remains below "
    "the 2021\u20132023 pre-period average of approximately $0.130 million. The "
    "post-Betsy productivity recovery is therefore best characterized as a "
    "partial recovery toward, rather than past, pre-restructuring "
    "productivity levels. As an author-constructed supplementary indicator, "
    "the Operational Leverage Index (OLI) reaches 195.3 in 2025 with 2023 "
    "set to 100 (see Appendix A.1). The OLI is equally weighted across "
    "Revenue Growth, Revenue per Employee, and Inverse Expense Ratio "
    "indices, and is sensitive to the choice of baseline year and weighting "
    "scheme. It is therefore retained as a compact visualization of "
    "multiple normalized trends rather than as primary evidence."
)

H3("5.3.5 Benchmark Context with Rocket Mortgage")

P(
    "Rocket Companies is included in this analysis as a directional "
    "industry benchmark, not as a control group. The two firms differ "
    "substantially in scale (Rocket\u2019s 2024 revenue exceeded $5 billion, "
    "compared with Better\u2019s approximately $108.5 million), in business "
    "mix, and in their 2025 strategic context: Rocket\u2019s 2025 financial "
    "results reflect the Mr. Cooper acquisition, which materially distorts "
    "per-employee comparisons (Rocket Companies, 2024, 2025). With those "
    "caveats, the benchmark exercise indicates that across 2023\u20132025 both "
    "originators experienced cyclical pressure followed by partial recovery. "
    "Better\u2019s relative improvement in expense ratio (4.79 \u2192 2.01) is more "
    "pronounced than Rocket\u2019s (1.106 \u2192 1.032), and Better\u2019s relative "
    "movement in profit margin and revenue per employee is also more "
    "favorable. These descriptive differences are consistent with the "
    "interpretation that Better\u2019s post-transition trajectory is not "
    "explained solely by macroeconomic recovery in the U.S. mortgage market, "
    "since Rocket experienced the same macroeconomic environment yet "
    "displayed a smaller relative improvement."
)

P(
    "This descriptive benchmark comparison is not equivalent to a formal "
    "econometric counterfactual. The descriptive benchmark comparison "
    "with Rocket Mortgage contextualizes Better.com\u2019s trajectory against "
    "broad market recovery effects rather than controlling for them. It is "
    "therefore reported in support of the descriptive operational-leverage "
    "interpretation, alongside the explicit acknowledgement that scale, "
    "business model, and acquisition-context differences preclude a "
    "treatment-effect interpretation."
)

H3("5.3.6 Financial Interpretation Boundaries")

P(
    "Several boundary conditions limit the strength of the financial "
    "evidence presented in Sections 5.3.1\u20135.3.5. First, the analysis "
    "rests on annual financial data, but Betsy launched on October 17, "
    "2024 (Better Home & Finance Holding Company, 2024a). Fiscal year 2024 "
    "therefore blends approximately ten months of pre-deployment performance "
    "with only two months of post-deployment performance, which is why the "
    "2024-to-2025 window is used as the primary post-transition comparison "
    "and why the 2023-to-2025 window is reported only as supporting "
    "recovery-from-trough evidence. Second, multiple non-AI factors "
    "operated concurrently with Betsy\u2019s deployment, including the "
    "ongoing maturation of Tinman, the launch of the NEO channel in early "
    "2025, broader cost-restructuring initiatives, and partial recovery "
    "of the U.S. mortgage market (Better Home & Finance Holding Company, "
    "2025a, 2026a; Mortgage Bankers Association, 2024). Third, the "
    "company-reported net-loss and revenue figures are audited, but the "
    "OLI is author-constructed and supplementary. Taken together, these "
    "boundary conditions imply that the financial findings should be "
    "interpreted as descriptively consistent with AI-enabled operational "
    "leverage, while leaving the analysis unable to isolate Betsy\u2019s "
    "specific contribution from the broader Betsy/Tinman operating model."
)

# ---------------------------------------------------------------------------
# 5.4 RQ2 - Customer Experience and NLP Evidence
# ---------------------------------------------------------------------------

H2("5.4 RQ2: Customer Experience and NLP Evidence")

P(
    "This section addresses the customer-experience dimension of RQ2 by "
    "integrating two complementary signals: management-reported Net "
    "Promoter Scores (NPS) and the independent NLP analysis of Trustpilot "
    "reviews. The argument developed across the four subsections is not "
    "that the NLP evidence proves a deterioration in customer experience "
    "post-Betsy, but that customer-experience evidence is mixed and "
    "methodologically valuable because it complicates the management-"
    "reported NPS narrative. The pre-Betsy and post-Betsy windows are "
    "defined relative to Betsy\u2019s October 17, 2024 launch date (Better "
    "Home & Finance Holding Company, 2024a), with 1,746 reviews falling "
    "into the pre-period and 169 into the post-period."
)

H3("5.4.1 Management-Reported NPS as Medium-Reliability Evidence")

P(
    "Better.com\u2019s public materials report that NPS improved from "
    "approximately 39 to approximately 64 over the relevant post-deployment "
    "period (Better Home & Finance Holding Company, 2025b, 2026a). At face "
    "value, this is a substantial improvement in an industry-standard "
    "customer-experience indicator and would be consistent with an "
    "interpretation that Betsy supports a faster, more responsive customer "
    "experience. NPS is therefore included as an evidentiary input. "
    "However, the NPS figure is management-reported and unaudited, and the "
    "underlying sampling, weighting, and survey-instrument details are not "
    "fully disclosed. Following the evidence-hierarchy logic established "
    "in Chapter 3, NPS is treated as medium-reliability evidence: it is a "
    "meaningful internal customer-experience signal, but it should not "
    "dominate the customer-experience interpretation, particularly when "
    "an independent customer-discourse stream is available."
)

H3("5.4.2 Trustpilot Ratings and VADER Sentiment as Independent Customer Discourse")

P(
    "The independent customer-discourse stream is a corpus of 1,915 cleaned "
    "Trustpilot reviews of Better.com (1,746 pre-Betsy; 169 post-Betsy). "
    "Two complementary signals are computed on the full clean dataset: "
    "the average star rating and the lexicon-based VADER compound "
    "sentiment score (Hutto & Gilbert, 2014). On both signals, the "
    "post-Betsy period is associated with statistically significant "
    "declines. Mean star ratings fall from approximately 4.2806 to 3.8580, "
    "a decline of approximately 0.4226 points, with Welch\u2019s t-test "
    "p = 0.0019 and Cohen\u2019s d = 0.2899 (small-to-medium effect). VADER "
    "sentiment falls from approximately 0.6416 to 0.5151, a decline of "
    "approximately 0.1265, with p = 0.0057 and d = 0.2505."
)

# Figure 5.4 - Sentiment by period
FIG(
    os.path.join(FIG_NLP, "fig_nlp_sentiment_by_period_clean.png"),
    "Figure 5.4. Distribution of Trustpilot VADER compound sentiment scores "
    "for Better.com pre-Betsy (n = 1,746) and post-Betsy (n = 169). The "
    "post-Betsy distribution is shifted modestly downward and exhibits a "
    "longer negative tail. Source: Author\u2019s analysis of 1,915 cleaned "
    "Trustpilot reviews (Apify scrape; collected on Better.com\u2019s public "
    "Trustpilot profile). Welch\u2019s t-test p = 0.0057, Cohen\u2019s d = 0.2505. "
    "The post-Betsy sample is substantially smaller than the pre-Betsy "
    "sample; the figure represents unsolicited Trustpilot review discourse "
    "rather than the full Better.com customer population.",
)

P(
    "Three interpretive moves are required to read this evidence "
    "responsibly. First, the post-Betsy sample is approximately ten times "
    "smaller than the pre-Betsy sample (169 vs. 1,746), meaning the "
    "post-period estimate is more sensitive to a small number of unusually "
    "negative reviews. Second, Trustpilot reviewers are self-selected and "
    "prone to extreme-response bias (Filieri, 2015; Luca, 2016), and the "
    "platform is therefore not representative of the full Better.com "
    "customer base. Third, the evidence is associative rather than causal: "
    "many factors changed concurrently with Betsy\u2019s launch, including "
    "product mix, market conditions, and broader service operations, and "
    "the keyword analysis cannot determine whether individual reviewers "
    "interacted with Betsy. With those caveats, the most defensible "
    "interpretation is that independent customer discourse on Trustpilot "
    "does not mirror the management-reported NPS improvement and instead "
    "displays a small-to-medium statistically significant decline in both "
    "ratings and sentiment over the post-Betsy window."
)

H3("5.4.3 Keyword Prevalence and the AI-Invisibility Finding")

P(
    "A targeted keyword-prevalence analysis was conducted across seven "
    "thematic groups on the full 1,915-review clean dataset. Reviews were "
    "binary-coded for the presence of group-specific keywords using strict "
    "word-boundary regular expressions, and Pearson\u2019s chi-squared tests "
    "evaluated whether prevalence shifted between the pre- and post-Betsy "
    "periods. A Bonferroni correction was applied to mitigate Type I error "
    "risk from multiple comparisons. The Cost/Rates group declined from "
    "28.92 percent to 16.57 percent (\u22128.98 pp before rounding to 12.35 "
    "percentage points; p = 0.0009), surviving strict Bonferroni correction. "
    "The Digital/Process group declined from 32.65 percent to 23.67 percent "
    "(\u22128.98 pp; p = 0.0211), significant under the unadjusted alpha but "
    "not after Bonferroni. Speed/Efficiency declined from 52.00 percent to "
    "44.38 percent (p = 0.0698), suggestive but not statistically "
    "significant. Human Support, Communication, and Delay/Friction did not "
    "change significantly."
)

# Figure 5.5 - Keyword group prevalence
FIG(
    os.path.join(FIG_NLP, "fig_nlp_keyword_prevalence_clean.png"),
    "Figure 5.5. Prevalence of seven keyword groups in pre-Betsy "
    "(n = 1,746) and post-Betsy (n = 169) Trustpilot reviews. Only the "
    "Cost/Rates decline survives strict Bonferroni correction (p = 0.0009). "
    "Explicit AI/Automation mentions remain near 1 percent in both periods "
    "(Fisher exact p = 1.0000). Source: Author\u2019s analysis of 1,915 "
    "cleaned Trustpilot reviews using strict word-boundary regular "
    "expressions and Pearson\u2019s chi-squared tests with Bonferroni "
    "correction.",
)

P(
    "The most thesis-relevant keyword result is the Explicit AI/Automation "
    "group. The prevalence of explicit AI mentions is essentially "
    "unchanged across the two periods (1.15 percent pre-Betsy versus 1.18 "
    "percent post-Betsy; Fisher exact p = 1.0000). This near-zero "
    "difference is interpreted in this thesis as the AI-invisibility "
    "finding: the near-zero explicit AI mention rate suggests that "
    "Trustpilot reviewers do not usually frame their mortgage experience "
    "in AI-specific terms. This supports the interpretation of voice AI as "
    "invisible infrastructure (Mogaji & Nguyen, 2021), although the "
    "keyword analysis cannot determine whether individual reviewers "
    "interacted with Betsy. The finding is therefore a discourse "
    "observation about what customers write about, rather than a "
    "behavioral claim about who interacted with the AI agent."
)

P(
    "The Cost/Rates result deserves a complementary interpretive note. "
    "A decline in the prevalence of cost- and rate-related keywords "
    "post-Betsy does not necessarily mean that Better\u2019s pricing improved "
    "in customers\u2019 experience. It indicates that the salience of "
    "cost- and rate-related language in the Trustpilot corpus declined. "
    "The most plausible compositional explanation is the broader "
    "interest-rate environment and a partial shift in product mix, "
    "rather than a direct effect of voice AI on price perception. The "
    "Digital/Process result is treated as suggestive only, given that it "
    "fails strict Bonferroni correction; this is consistent with the "
    "approach in Wang et al. (2023) and Zhang et al. (2023), which "
    "emphasize cautious interpretation of NLP signals when sample sizes "
    "differ markedly between groups."
)

H3("5.4.4 Topic Modeling and Product-Mix Interpretation")

P(
    "BERTopic neural topic modeling (Grootendorst, 2022) was applied to a "
    "1,804-document subset of the cleaned Trustpilot corpus, after "
    "filtering CEO- and 2021-controversy-related reviews that focused on "
    "the company\u2019s leadership rather than its mortgage product. The "
    "pipeline produced 44 topics; 1,198 documents were successfully "
    "clustered and 606 were assigned to the HDBSCAN noise class. The "
    "clustered non-outlier subsample used for the topic-share comparison "
    "consists of 1,077 pre-Betsy and 121 post-Betsy reviews. Because the "
    "post-Betsy clustered non-outlier sample is small, topic-share "
    "differences are interpreted as descriptive signals rather than as "
    "statistically tested population-level changes."
)

# Figure 5.6 - Top topics
FIG(
    os.path.join(FIG_NLP, "fig_nlp_topic_shares_clean.png"),
    "Figure 5.6. Selected topic-share changes between pre-Betsy "
    "(n = 1,077) and post-Betsy (n = 121) clustered non-outlier Trustpilot "
    "reviews. The figure is presented as a descriptive overview of the "
    "topical composition shift, not as a statistical test. Refinance "
    "topics decline by approximately 7.66 percentage points combined; "
    "HELOC Complaints, Smooth Transaction, Responsive Service, and "
    "Closing Timeline rise by between 3.55 and 4.22 percentage points. "
    "Source: Author\u2019s analysis using BERTopic with all-MiniLM-L6-v2 "
    "embeddings, UMAP, and HDBSCAN.",
)

P(
    "Read descriptively, several patterns are visible. The combined share "
    "of refinance-oriented topics declines by approximately 7.66 "
    "percentage points, while the share of HELOC-related complaints rises "
    "by approximately 4.22 percentage points and the share of Home "
    "Purchase topics declines by approximately 3.48 percentage points. "
    "The Smooth Transaction (+4.03 pp), Responsive Service (+3.93 pp), "
    "and Closing Timeline (+3.55 pp) topics also rise modestly. The most "
    "plausible interpretation of the refinance decline is the macro-level "
    "interest-rate environment, which suppressed refinancing demand "
    "across the U.S. mortgage industry rather than reflecting a Betsy "
    "effect (Mortgage Bankers Association, 2024). The rise in HELOC "
    "complaints is most plausibly read as a product-mix signal, given "
    "Better\u2019s strategic emphasis on the HELOC product over the relevant "
    "period. The simultaneous rise of Smooth Transaction and Responsive "
    "Service shares is directionally consistent with the operational-"
    "efficiency interpretation, but cannot be associated solely with Betsy "
    "given the small post-period sample, the absence of controls, and "
    "concurrent product- and process-mix changes."
)

H3("5.4.5 Customer-Experience Boundaries: Synthesis for the NLP Stream")

P(
    "Synthesizing the customer-experience evidence, the NLP stream answers "
    "RQ2\u2019s customer-experience dimension as follows. Independent "
    "Trustpilot discourse displays a small-to-medium, statistically "
    "significant decline in star ratings and VADER sentiment across the "
    "post-Betsy window, while management-reported NPS rises substantially. "
    "Customers rarely frame their experience in explicit AI terms, which "
    "is consistent with the interpretation of voice AI as invisible "
    "infrastructure but is silent on individual exposure to Betsy. Topic "
    "shifts are best interpreted as descriptive evidence of product-mix "
    "and macro-environment effects rather than as evidence of Betsy\u2019s "
    "direct effect on customer experience. The customer-experience stream "
    "therefore complicates rather than confirms the financial-operational "
    "interpretation, and this complication is itself analytically valuable."
)

# ---------------------------------------------------------------------------
# 5.5 Triangulated Interpretation
# ---------------------------------------------------------------------------

H2("5.5 Triangulated Interpretation")

P(
    "Taken together, the three evidence streams support a multidimensional "
    "interpretation of voice AI business value at Better.com. The "
    "qualitative case evidence shows that Betsy is not a stand-alone "
    "chatbot but a voice interface embedded in Better\u2019s Tinman-based "
    "mortgage operating platform, with ElevenLabs as voice infrastructure "
    "and licensed humans retaining authority over consequential decisions. "
    "The financial evidence shows that the post-transition period is "
    "associated with revenue recovery, expense-intensity improvement, "
    "narrowing losses, and partial labor-productivity recovery. The "
    "customer-experience evidence is mixed: management-reported NPS "
    "improves materially, while independent Trustpilot star ratings and "
    "sentiment decline significantly, and explicit AI mentions remain "
    "near-zero in both periods."
)

P(
    "The strongest evidence is for the operational-efficiency and "
    "financial-trajectory dimensions of business value. The expense-ratio "
    "improvement from 2.89 to 2.01 between 2024 and 2025, achieved with "
    "approximately 52 percent revenue growth against approximately 5 "
    "percent expense growth, is the textbook descriptive signature of "
    "operational leverage and is consistent with the AI-enabled "
    "operational-leverage hypothesis examined in this thesis. The "
    "evidence for productivity and scalability is moderately strong: RPE "
    "improves substantially in 2025 but remains below the 2021\u20132023 "
    "pre-period average of approximately $0.130 million, indicating "
    "partial rather than full recovery. The evidence for customer-"
    "experience improvement is weak and divergent: NPS supports an "
    "improvement narrative, but the Trustpilot signals run in the "
    "opposite direction. The case evidence on RQ1 supplies the integration "
    "mechanism that makes the financial signature plausible: a "
    "Tinman-mediated workflow into which a voice LLM agent has been "
    "embedded with auditable human escalation."
)

P(
    "The divergence between the financial-operational and customer-discourse "
    "streams is methodologically valuable rather than analytically "
    "embarrassing. It indicates that different measurement instruments "
    "capture different facets of customer experience and business value. "
    "Management-reported NPS captures customers\u2019 internal evaluation in a "
    "structured survey context. Trustpilot ratings and VADER sentiment "
    "capture self-selected, unsolicited public discourse from a smaller, "
    "more vocal subset of customers. Topic and keyword analyses capture "
    "the linguistic salience of particular themes in that public "
    "discourse, which is sensitive to product-mix and macro-environment "
    "effects. The most defensible synthesis is therefore that voice AI in "
    "the Better.com case is associated with measurable improvements in "
    "operational efficiency and financial trajectory, while its effects on "
    "customer experience are uneven, mediated, and possibly under-detected "
    "by any single instrument. This is consistent with Aggarwal et al.\u2019s "
    "(2025) argument that voice AI value creation occurs across both "
    "B2B and B2C surfaces in ways that are not captured by a single "
    "metric, and with the broader literature on AI in customer-facing "
    "financial services (Hentzen et al., 2022; Mogaji & Nguyen, 2021)."
)

P(
    "The triangulated interpretation also clarifies the limits of attribution. "
    "Even where the financial evidence is strong, it cannot isolate Betsy\u2019s "
    "specific contribution from the broader Betsy/Tinman AI operating "
    "model. The case evidence indicates that Betsy operates as one "
    "component of an integrated technology stack and operating model, "
    "alongside Tinman, ElevenLabs, the NEO channel, restructuring, and "
    "broader market dynamics. The thesis therefore does not claim that "
    "Betsy alone caused Better.com\u2019s post-transition financial trajectory. "
    "It claims that the post-transition trajectory is consistent with "
    "AI-enabled operational leverage when interpreted alongside the "
    "qualitative case evidence on Betsy\u2019s deployment and Tinman\u2019s "
    "architectural role, and that customer-experience effects appear to "
    "be more uneven than the financial signature would suggest."
)

# ---------------------------------------------------------------------------
# 5.6 Limitations and boundary conditions
# ---------------------------------------------------------------------------

H2("5.6 Limitations and Boundary Conditions of the Empirical Findings")

P(
    "Five boundary conditions should be kept in mind when reading the "
    "empirical findings. First, the study is a single-case design and "
    "therefore does not support statistical generalization to all U.S. "
    "mortgage originators or to financial services more broadly; Better\u2019s "
    "digital-first architecture and Tinman platform are unusually "
    "favorable conditions for voice LLM integration. Second, the financial "
    "data are reported at annual granularity, which limits the temporal "
    "resolution of the post-Betsy window; the 2024 fiscal year blends "
    "approximately ten months of pre-deployment performance with two "
    "months of post-deployment performance. Third, the NLP corpus is drawn "
    "from a single review platform, is unsolicited, and is dominated by "
    "self-selected reviewers; the pre/post imbalance (1,746 vs. 169) "
    "further limits the precision of the post-period estimates. Fourth, "
    "Rocket Companies is included as a directional industry benchmark "
    "rather than a counterfactual; differences in scale, business model, "
    "and the 2025 Mr. Cooper acquisition preclude treatment-effect "
    "interpretation. Fifth, vendor materials from ElevenLabs are subject "
    "to an obvious incentive structure and are therefore used to "
    "characterize the integration mechanism rather than to substantiate "
    "magnitude claims about cost savings or revenue uplift."
)

P(
    "These limitations do not invalidate the findings, but they bound the "
    "kind of claim the thesis can make. The strongest defensible claims "
    "are descriptive and triangulated: the post-deployment period is "
    "associated with revenue recovery, expense-intensity improvement, "
    "narrowing losses, and partial productivity recovery; customer "
    "discourse on Trustpilot is mixed and includes a statistically "
    "significant rating and sentiment decline; explicit AI mentions are "
    "rare; and the case-study evidence supplies the workflow-integration "
    "mechanism through which a voice LLM agent could plausibly contribute "
    "to operational leverage."
)

# ---------------------------------------------------------------------------
# 5.7 Summary of empirical findings
# ---------------------------------------------------------------------------

H2("5.7 Summary of Empirical Findings")

P(
    "Chapter 5 has integrated qualitative case-study evidence, longitudinal "
    "financial evidence, and customer-discourse evidence into a single "
    "empirical argument addressing the two research questions of the "
    "thesis."
)

P(
    "RQ1 (integration). Voice LLM agents become operationally meaningful "
    "in mortgage origination when they are embedded in a process-aware "
    "operating platform that exposes structured workflow state, when they "
    "are scoped to a communication-coordination role rather than a "
    "regulated decision-making role, when they are supported by low-"
    "latency voice infrastructure, and when they are governed by an "
    "explicit human-AI collaboration model with auditable escalation. In "
    "the Better.com case, Tinman supplies the architectural prerequisite, "
    "Betsy operationalizes the conversational coordination layer, "
    "ElevenLabs supplies the voice infrastructure, and the regulated "
    "task structure of mortgage origination supplies the human-in-the-loop "
    "boundary."
)

P(
    "RQ2 (business value). The empirical evidence is consistent with "
    "AI-enabled operational leverage in the financial-operational "
    "dimension and is mixed in the customer-experience dimension. "
    "Concretely, the post-transition period (2024\u20132025) is associated "
    "with approximately 52.0 percent revenue growth and only "
    "approximately 5.3 percent expense growth, an expense-ratio decline "
    "from 2.89 to 2.01, an approximately 19.6 percent narrowing of net "
    "loss, and an approximately 42.9 percent improvement in revenue per "
    "employee \u2014 although the company remained unprofitable in 2025 and "
    "RPE remained below the pre-period average. Customer-experience "
    "evidence is mixed: management-reported NPS rises from 39 to 64, "
    "while independent Trustpilot star ratings decline from 4.2806 to "
    "3.8580 (p = 0.0019; d = 0.2899) and VADER sentiment declines from "
    "0.6416 to 0.5151 (p = 0.0057; d = 0.2505). Explicit AI mentions "
    "remain near 1 percent in both periods. Topic shifts are descriptive "
    "and most plausibly reflect product-mix and macro-environment effects "
    "rather than direct Betsy effects."
)

P(
    "Synthesized across the three streams, the central empirical finding "
    "of the thesis is that voice AI in financial services should be "
    "understood as creating multidimensional, mediated business value "
    "rather than uniformly directional improvement. The financial "
    "signature is consistent with AI-enabled operational leverage but "
    "cannot isolate Betsy\u2019s specific contribution; the customer "
    "discourse is mixed in ways that complicate, but do not invalidate, "
    "the management-reported NPS narrative; and the case-study evidence "
    "supplies the workflow-integration mechanism through which the "
    "financial signature becomes plausible. Chapter 6 develops the "
    "managerial and theoretical implications of these findings, including "
    "the design pattern that the Better.com case suggests for voice LLM "
    "integration in regulated, communication-intensive financial-services "
    "workflows."
)

# ---------------------------------------------------------------------------
# Pipeline driver
# ---------------------------------------------------------------------------


# Bibliography entries to add after the existing references list. Each entry
# is appended verbatim as a new "Normal" paragraph at the end of the
# References section. Entries that already exist in the bibliography (per the
# audit performed during the cleanup pass) are NOT re-added here.
NEW_REFERENCES: list[str] = [
    "Better.com. (n.d.). Our story. Better.com. Retrieved May 6, 2026, from "
    "https://better.com",
    "Rocket Companies, Inc. (2024). Form 10-K: Annual report for the fiscal "
    "year ended December 31, 2024. U.S. Securities and Exchange Commission.",
    "Rocket Companies, Inc. (2025). Form 10-K: Annual report for the fiscal "
    "year ended December 31, 2025. U.S. Securities and Exchange Commission.",
]


def find_paragraph_with_text(doc, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    return None


def chapter5_already_present(doc) -> bool:
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and para.text.strip().startswith(
            "5. Empirical Findings"
        ):
            return True
    return False


def existing_reference_text(doc, needle: str) -> bool:
    """Return True if any paragraph in the References section already
    contains the needle string."""
    in_refs = False
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and para.text.strip() == "References":
            in_refs = True
            continue
        if not in_refs:
            continue
        if needle in para.text:
            return True
    return False


def add_reference_paragraph(doc, anchor_para, text: str) -> None:
    """Append a new reference paragraph at the end of the References list,
    using the same paragraph style as the anchor (typically the last
    existing reference)."""
    new_para = insert_paragraph_before(anchor_para, text, style_name="Normal")
    return new_para


def main() -> int:
    if not os.path.exists(DRAFT_PATH):
        print(f"Draft not found at {DRAFT_PATH}", file=sys.stderr)
        return 1

    doc = docx.Document(DRAFT_PATH)

    if chapter5_already_present(doc):
        print(
            "[skip] Chapter 5 ('5. Empirical Findings: ...') is already "
            "present in the draft; not inserting again."
        )
        return 0

    refs_para = find_paragraph_with_text(doc, "References")
    if refs_para is None:
        print("[error] Could not find a Heading 1 paragraph titled 'References'", file=sys.stderr)
        return 2

    # Insert each chapter element directly before the References heading.
    style_for = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "p": "Normal",
    }

    for kind, payload in CONTENT:
        if kind == "fig":
            image_path, caption = payload
            if not os.path.exists(image_path):
                print(f"[warn] Figure not found: {image_path}", file=sys.stderr)
                continue
            insert_picture_before(refs_para, image_path, width_in=6.0)
            cap_para = insert_paragraph_before(refs_para, caption, style_name="Caption")
        else:
            insert_paragraph_before(refs_para, str(payload), style_name=style_for[kind])

    # Append new bibliographic entries (after the last existing reference,
    # which is the last paragraph before the end-of-document or any future
    # appendix). Implementation: simply append to end of body.
    for ref_text in NEW_REFERENCES:
        # De-duplicate against existing references by matching on a unique
        # substring such as the year + leading author surname.
        key = ref_text.split("(")[0].strip()
        if existing_reference_text(doc, key):
            print(f"[skip] Reference already present: {key}")
            continue
        para = doc.add_paragraph(ref_text)
        para.style = doc.styles["Normal"]

    doc.save(DRAFT_PATH)
    print(f"Inserted Chapter 5 into {DRAFT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
