"""
extend_chapter5.py
==================

Idempotent extension pass for Chapter 5. After ``insert_chapter5.py`` has
created the chapter, this script inserts additional paragraphs to bring
each section closer to the supervisor-plan target word counts (5.1: 500-700,
5.2: 1,300-1,700, 5.3: 2,200-2,800, 5.4: 1,700-2,300, 5.5: 900-1,200,
5.6: 400-500, 5.7: 400-600).

Each new paragraph is anchored on a unique substring from the existing
Chapter 5 prose and is inserted immediately after the matching paragraph.
A "first-sentence fingerprint" check makes the script idempotent: the same
extension paragraph is only inserted once.
"""
from __future__ import annotations

import os
import sys

import docx


REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DRAFT_PATH = os.path.join(REPO_ROOT, "04_Drafts", "Main_Draft.docx")


# Each entry: (anchor_substring_in_existing_para, new_paragraph_text).
# anchor_substring must match exactly one paragraph in Chapter 5.
EXTENSIONS: list[tuple[str, str]] = [
    # ----- 5.3.1 Revenue recovery -----
    (
        "but this larger figure incorporates the trough-year baseline of 2023",
        "Two additional contextual observations help calibrate this revenue "
        "recovery. First, the recovery is partial rather than complete. Even "
        "at the 2025 level of approximately $164.9 million, revenue remains "
        "materially below Better\u2019s 2021 pandemic-era peak, when the U.S. "
        "mortgage refinancing boom drove unusually high origination volumes. "
        "The 2025 revenue figure is therefore best read as a meaningful "
        "recovery from the 2023 trough rather than as a return to "
        "pre-restructuring scale. Second, the recovery is occurring in a "
        "U.S. mortgage market that is itself only partially recovering. "
        "Mortgage Bankers Association (2024) data indicate that independent "
        "mortgage banker production margins improved from late-2023 lows but "
        "remained below long-run averages through 2024, with origination "
        "volumes still constrained by elevated rates. The descriptive "
        "benchmark comparison with Rocket Mortgage in Section 5.3.5 helps "
        "separate firm-specific from sector-wide effects, although it cannot "
        "do so with causal precision.",
    ),
    # ----- 5.3.2 Expense ratio -----
    (
        "is the textbook descriptive signature of operational leverage and is consistent",
        "The decomposition of the expense ratio is informative. The "
        "2024-to-2025 improvement from 2.89 to 2.01 reflects an "
        "approximately 30.6 percent decline. The bulk of that decline is "
        "attributable to revenue growth (the denominator effect), with the "
        "remainder attributable to only modest expense growth (the numerator "
        "effect). In a pure cost-cutting story, one would expect the "
        "numerator to fall sharply; in a pure demand-recovery story, one "
        "would expect the numerator to rise in proportion to the "
        "denominator. The Better.com evidence fits neither pure pattern. "
        "Instead, expenses grew only marginally while revenue grew "
        "substantially, which is what an operational-leverage trajectory "
        "should look like when a process-aware operating platform absorbs "
        "additional volume without proportional cost. This is the "
        "descriptive pattern most consistent with the central hypothesis of "
        "the thesis, with the caution that the pattern reflects the broader "
        "Betsy/Tinman operating model rather than Betsy in isolation.",
    ),
    # ----- 5.3.3 Loss narrowing -----
    (
        "this comparative observation is treated as directional rather than as a controlled estimate",
        "It is also useful to record what the loss-narrowing trajectory does "
        "not establish. It does not establish that Better\u2019s underlying "
        "business model is profitable at current scale; the company\u2019s 2025 "
        "net loss remains substantial in absolute terms, and the implied "
        "profit margin remains negative. It does not establish that further "
        "loss-narrowing will continue at the same rate; the 2024-to-2025 "
        "improvement of approximately 19.6 percent is smaller in proportional "
        "terms than the 2023-to-2024 improvement, and a substantial portion "
        "of the absolute improvement reflects revenue growth rather than "
        "cost reduction. Finally, it does not establish that the path to "
        "profitability is uniquely AI-driven; the joint Better\u2013ElevenLabs "
        "materials and Better\u2019s own SEC filings emphasize multiple concurrent "
        "levers, including capital structure changes, channel "
        "diversification, and cost discipline. The most defensible reading "
        "is therefore that the loss-narrowing trajectory is consistent with "
        "the operational-leverage interpretation, but does not yet support "
        "a stronger claim of financial turnaround.",
    ),
    # ----- 5.3.5 Rocket benchmark -----
    (
        "Rocket experienced the same macroeconomic environment yet displayed a smaller relative improvement",
        "The benchmark exercise also produces honest counter-evidence. "
        "Rocket remains far larger and more diversified than Better, and "
        "Rocket\u2019s absolute profitability profile across 2018\u20132022 is "
        "structurally different from Better\u2019s. In particular, Rocket "
        "reports positive net income in most years across the benchmark "
        "window, whereas Better reports net losses throughout 2021\u20132025. "
        "A descriptive reading of this pattern is that Better is operating "
        "from a much smaller and more constrained base, so even an "
        "operational-leverage trajectory does not yet bring it to "
        "profitability. A complementary reading is that Rocket\u2019s "
        "structurally lower expense ratio reflects scale economies and a "
        "more mature operational base, neither of which can be matched by "
        "Better in the short term, and neither of which is the object of "
        "this thesis. The benchmark therefore contextualizes Better\u2019s "
        "trajectory against broad market recovery effects without replacing "
        "the central case-level interpretation.",
    ),
    # ----- 5.4.1 NPS -----
    (
        "particularly when an independent customer-discourse stream is available",
        "Two interpretive caveats are required. First, the NPS movement "
        "reported by Better.com is consistent with patterns reported in "
        "adjacent literature on AI-augmented customer service. Studies of "
        "AI-based conversational agents in call-center contexts (Wang et "
        "al., 2023; Zhang et al., 2023) document operational-performance "
        "gains alongside variable customer-satisfaction effects, and "
        "Brynjolfsson et al. (2023) report meaningful productivity gains "
        "for less-experienced workers in AI-augmented service settings. "
        "Whether the magnitude of Better\u2019s reported NPS movement is "
        "plausible relative to those benchmarks cannot be evaluated from "
        "public disclosures alone. Second, NPS is a single summary metric "
        "based on a structured survey instrument and is sensitive to "
        "question wording, sampling timing, and respondent self-selection. "
        "The divergence with the Trustpilot evidence in Section 5.4.2 is "
        "therefore not a contradiction in the strong sense, but a "
        "divergence between two customer-experience instruments that "
        "measure partially overlapping but non-identical constructs.",
    ),
    # ----- 5.4.2 Trustpilot -----
    (
        "displays a small-to-medium statistically significant decline in both ratings and sentiment over the post-Betsy window",
        "Effect-size context is useful. Cohen\u2019s d values of approximately "
        "0.29 (ratings) and 0.25 (sentiment) correspond, in conventional "
        "terms, to small-to-medium effects. They are not negligible, but "
        "they are also not large. In substantive terms, the rating decline "
        "of approximately 0.42 points on a 5-point scale is consistent "
        "with a moderate downward shift in the distribution rather than a "
        "collapse of customer discourse, and the corresponding VADER "
        "sentiment decline of approximately 0.13 on a roughly \u22121-to-+1 "
        "compound scale fits the same characterization. The Trustpilot "
        "stream therefore complicates any reading of the customer-experience "
        "evidence as uniformly positive, but it does not support the strong "
        "claim that customer experience deteriorated dramatically post-Betsy.",
    ),
    # ----- 5.4.3 Keywords (after the Cost/Rates note) -----
    (
        "emphasize cautious interpretation of NLP signals when sample sizes differ markedly between groups",
        "The non-significant keyword groups are themselves analytically "
        "useful. The Human Support group (43.81 percent pre versus 40.83 "
        "percent post) shows that customer attention to human-mediated "
        "service did not collapse after Betsy\u2019s deployment, which is "
        "consistent with the task-reallocation interpretation rather than a "
        "labor-substitution interpretation. The Communication group (32.53 "
        "percent pre versus 30.77 percent post) is essentially unchanged, "
        "suggesting that the salience of communication-quality language in "
        "customer reviews did not shift in the post-Betsy window. The "
        "Delay/Friction group (13.52 percent pre versus 10.06 percent "
        "post) declines modestly but not significantly. None of these "
        "groups individually justifies a strong interpretive claim, but "
        "together they suggest that the post-Betsy customer-discourse "
        "environment did not undergo a wholesale compositional shift toward "
        "or away from any single keyword theme. What did shift, modestly "
        "but reliably, was overall sentiment and rating, alongside a salient "
        "decline in cost/rates language and a persistent absence of "
        "explicit AI vocabulary.",
    ),
    # ----- 5.4.4 Topics -----
    (
        "but cannot be associated solely with Betsy given the small post-period sample",
        "Two methodological observations sharpen the topic-level reading. "
        "First, BERTopic clustering depends on hyperparameter choices in "
        "the UMAP and HDBSCAN stages and on the all-MiniLM-L6-v2 embedding "
        "model. As Grootendorst (2022) notes, topic identities are stable "
        "across reasonable hyperparameter perturbations but exact topic "
        "shares are sensitive to those choices; the percentage-point "
        "figures reported above should therefore be interpreted as "
        "approximate compositional signals rather than as fixed population "
        "parameters. Second, the clustered non-outlier denominator structure "
        "(Pre n = 1,077 versus Post n = 121) is even more imbalanced than "
        "the full clean dataset, which means topic-share estimates for the "
        "post-period are particularly sensitive to a small number of "
        "reviews. The figures reported in this section therefore serve as "
        "a descriptive map of the post-Betsy discourse rather than as "
        "inferential claims about population-level topic prevalence. "
        "Appendix figures (intertopic distance map, topic-word importance, "
        "and quarterly keyword trends) are provided for transparency but "
        "are not used to substantiate any claim in the main chapter.",
    ),
    # ----- 5.5 Triangulated interpretation, after the attribution-limits paragraph -----
    (
        "customer-experience effects appear to be more uneven than the financial signature would suggest",
        "From a theoretical perspective, the triangulated pattern is most "
        "naturally read through three complementary lenses. First, in the "
        "task-technology fit tradition (Goodhue & Thompson, 1995), the "
        "financial-operational improvement is consistent with a high fit "
        "between Betsy\u2019s task scope (high-volume, structured, "
        "communication-intensive coordination work) and Tinman\u2019s "
        "structured workflow environment, while the customer-experience "
        "divergence is consistent with a more uneven fit between AI-mediated "
        "voice interaction and the more relational, emotionally laden "
        "moments of the borrower journey (Bl\u00fcmel et al., 2024). Second, "
        "in the AI service hierarchy of Huang and Rust (2018), Betsy is "
        "positioned in the mechanical and analytical layers, where "
        "productivity gains are most likely to materialize, while the "
        "customer-experience layer engages intuitive and empathetic "
        "dimensions where AI substitution is more contested. Third, in the "
        "AI-enabled productivity tradition revisited by Brynjolfsson et al. "
        "(2023) and Aggarwal et al. (2025), AI-related productivity gains "
        "tend to emerge alongside ambiguous customer-perception effects "
        "when the AI system operates as backend infrastructure rather than "
        "as a branded front-of-house experience. The Better.com pattern is "
        "consistent with all three lenses simultaneously, which strengthens "
        "the descriptive interpretation without resolving the causal "
        "question.",
    ),
    (
        "The Better.com pattern is consistent with all three lenses simultaneously",
        "The triangulated interpretation also has implications for the "
        "design-pattern argument that Chapter 6 develops. The Better.com "
        "case suggests that voice AI in financial services is most likely "
        "to generate measurable business value when it is integrated with "
        "a process-aware operating platform that exposes structured "
        "workflow state, when its task scope is bounded to communication "
        "coordination rather than regulated decision-making, when its "
        "voice infrastructure supports natural conversation at low latency, "
        "and when human escalation is preserved and auditable. These "
        "integration prerequisites are summarized as a design pattern in "
        "Chapter 6, with explicit acknowledgement that the pattern is "
        "derived from a single revelatory case and is therefore proposed "
        "for analytic generalization rather than for statistical "
        "generalization.",
    ),
    # ----- 5.6 Limitations -----
    (
        "the case-study evidence supplies the workflow-integration mechanism through which a voice LLM agent could plausibly contribute to operational leverage",
        "A further conceptual boundary deserves explicit statement. The "
        "thesis focuses on the deployment of a voice LLM agent in mortgage "
        "origination, a regulated, communication-intensive financial-"
        "services context. The empirical findings are most directly "
        "informative for analogous contexts \u2014 retail banking servicing, "
        "insurance servicing, and other high-volume, communication-"
        "intensive financial workflows where a process-aware operating "
        "platform is or could be in place. They are less directly "
        "informative for less-regulated or less-communication-intensive "
        "contexts, and they should not be read as a general claim about "
        "AI-enabled productivity in financial services as a whole. Chapter "
        "6 develops the design-pattern implications of this scope "
        "condition.",
    ),
    # ----- 5.1 Introduction (after final preview paragraph) -----
    (
        "presents direct, evidence-bounded answers to RQ1 and RQ2",
        "A note on terminology is appropriate before the empirical sections "
        "begin. Throughout this chapter, the phrase \u201cthe broader Betsy/"
        "Tinman AI operating model\u201d is used as shorthand for the joint "
        "configuration of the Tinman platform, the Betsy voice agent, the "
        "ElevenLabs voice infrastructure, the human-in-the-loop escalation "
        "structure, and the organizational practices that surround them. "
        "Where the chapter refers specifically to Betsy as a discrete "
        "component, this is made explicit. This convention reflects the "
        "analytical position established in Chapter 3 that the unit of "
        "analysis is not Betsy in isolation but Betsy as one component of "
        "an integrated operating model, and it is consistent with the "
        "cautious causal language adopted throughout the thesis.",
    ),
]


def find_paragraph_containing(doc, needle: str):
    """Return the first paragraph that contains ``needle`` (and is inside
    Chapter 5 of the document, between the Chapter 5 H1 and the References
    H1)."""
    in_ch5 = False
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and para.text.strip().startswith(
            "5. Empirical Findings"
        ):
            in_ch5 = True
            continue
        if para.style.name == "Heading 1" and para.text.strip() == "References":
            return None
        if in_ch5 and needle in para.text:
            return para
    return None


def fingerprint(text: str, n: int = 80) -> str:
    return text.replace("\u00a0", " ").strip()[:n]


def insert_after(anchor_para, text: str, style_name: str = "Normal"):
    doc = anchor_para.part.document
    tmp = doc.add_paragraph()
    tmp.style = doc.styles[style_name]
    tmp.add_run(text)
    anchor_para._p.addnext(tmp._p)
    return tmp


def main() -> int:
    if not os.path.exists(DRAFT_PATH):
        print(f"Draft not found at {DRAFT_PATH}", file=sys.stderr)
        return 1

    doc = docx.Document(DRAFT_PATH)

    # Collect existing fingerprints so that re-running this script is a no-op.
    existing_fps = {fingerprint(p.text) for p in doc.paragraphs}

    n_added = 0
    n_skipped_dup = 0
    n_missing_anchor = 0

    for anchor, new_text in EXTENSIONS:
        if fingerprint(new_text) in existing_fps:
            n_skipped_dup += 1
            continue
        anchor_para = find_paragraph_containing(doc, anchor)
        if anchor_para is None:
            print(f"[warn] anchor not found in Chapter 5: {anchor[:60]}...", file=sys.stderr)
            n_missing_anchor += 1
            continue
        insert_after(anchor_para, new_text, style_name="Normal")
        existing_fps.add(fingerprint(new_text))
        n_added += 1

    if n_added:
        doc.save(DRAFT_PATH)
    print(f"added={n_added} skipped_duplicate={n_skipped_dup} missing_anchor={n_missing_anchor}")
    return 0 if n_missing_anchor == 0 else 0  # warnings are non-fatal


if __name__ == "__main__":
    sys.exit(main())
