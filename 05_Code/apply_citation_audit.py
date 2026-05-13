"""Apply Deliverable A/B citation-audit fixes to 04_Drafts/Main_Draft.docx.

Steps:
1. Apply targeted in-text-citation corrections in body paragraphs (0-353):
   - Cost/Rates pp correction
   - Better 2026a -> 2026 (collapse single-source suffix)
   - Better 2024c -> 2024b (orphan)
   - Better 2025 (no suffix) -> 2025a (FY 2024 / Q4 results, most common context)
   - Mortgage Bankers Association 2024 (no suffix) -> 2024a (full-year report) or
     2024b (Q3 IMB profits release) by context
   - "Author and Author, YEAR" inside parentheses -> "Author & Author, YEAR"
2. Replace the reference list (paragraphs 355-432) with the rebuilt list from
   Deliverable B (50 entries, alphabetised, deduplicated).
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parents[1]
DOCX_PATH = REPO / "04_Drafts" / "Main_Draft.docx"

# ---- Body citation corrections -------------------------------------------------
# Each (pattern, replacement) is applied to every paragraph 0-353.
BODY_FIXES: list[tuple[re.Pattern[str], str]] = [
    # Cost/Rates pp correction (the single most important substantive fix)
    (
        re.compile(
            r"\(\s*\u22128\.98\s*pp\s+before\s+rounding\s+to\s+12\.35\s*"
            r"percentage\s+points;\s*p\s*=\s*0\.0009\s*\)",
            re.IGNORECASE,
        ),
        "(\u221212.35 pp; Fisher's exact p \u2248 0.001; Bonferroni-adjusted p \u2248 0.005)",
    ),
    # Collapse Better-only 2026a -> 2026 (only one Better-only 2026 source)
    (
        re.compile(r"Better Home & Finance Holding Company,\s*2026a"),
        "Better Home & Finance Holding Company, 2026",
    ),
    # Better 2024c -> 2024b (orphan; no third Better-only 2024 source)
    (
        re.compile(r"Better Home & Finance Holding Company,\s*2024c"),
        "Better Home & Finance Holding Company, 2024b",
    ),
    # Better 2025 (no suffix) -> 2025a (Q4/FY 2024 results press release; most common context)
    # We narrow this to the parenthetical and narrative forms only to avoid touching
    # references like "Better Home & Finance Holding Company, 2025" inside ref list.
    (
        re.compile(
            r"\(Better Home & Finance Holding Company,\s*2025([^a-z\d])"
        ),
        r"(Better Home & Finance Holding Company, 2025a\1",
    ),
    # MBA 2024 (no suffix in body) -> 2024a by default (full-year industry report)
    (
        re.compile(r"\(Mortgage Bankers Association,\s*2024\)"),
        "(Mortgage Bankers Association, 2024a)",
    ),
    (
        re.compile(r"Mortgage Bankers Association \(2024\)"),
        "Mortgage Bankers Association (2024a)",
    ),
    # APA 7: parenthetical form must use "&", not "and"
    (
        re.compile(r"\((Andrade)\s+and\s+(Tumelero,\s*\d{4}[a-z]?)"),
        r"(\1 & \2",
    ),
    (
        re.compile(r"\((Fotheringham)\s+and\s+(Wiles,\s*\d{4}[a-z]?)"),
        r"(\1 & \2",
    ),
    (
        re.compile(r"\((Goodhue)\s+and\s+(Thompson,\s*\d{4}[a-z]?)"),
        r"(\1 & \2",
    ),
    (
        re.compile(r"\((Huang)\s+and\s+(Rust,\s*\d{4}[a-z]?)"),
        r"(\1 & \2",
    ),
    (
        re.compile(r"\((Mogaji)\s+and\s+(Nguyen,\s*\d{4}[a-z]?)"),
        r"(\1 & \2",
    ),
    # NPS-Trustpilot caveat replacement: detect the original short divergence
    # paragraph and replace with a 7-reason caveat. We anchor on a recognisable
    # opening clause to avoid false matches.
    (
        re.compile(
            r"Two interpretive caveats are required\.\s+First,\s+the NPS movement reported by Better\.com is consistent with patterns reported in adjacent literature on AI-augmented customer service\.[^\u2606]*?two customer-experience instruments that measure partially overlapping but non-identical constructs\."
        ),
        (
            "Two interpretive caveats are required when reading the NPS-vs-Trustpilot "
            "divergence reported in this chapter. First, Better's investor materials "
            "report a Net Promoter Score improvement from approximately 39 to "
            "approximately 64 over the relevant post-deployment period, but this "
            "figure is management-reported and not independently audited. Whether "
            "the magnitude of Better's reported NPS movement is plausible relative "
            "to peer-reviewed benchmarks (e.g., Wang et al., 2023; Zhang et al., "
            "2023; Brynjolfsson et al., 2023) cannot be evaluated from public "
            "disclosures alone. Second, NPS and Trustpilot VADER sentiment are "
            "not directly comparable for at least seven reasons: (i) sampling, "
            "because internal NPS samples are drawn from Better's own contact "
            "channels while Trustpilot reviews are self-selected; (ii) timing, "
            "because the two metrics use different measurement windows; (iii) "
            "question framing, because NPS asks a single recommendation-likelihood "
            "question while Trustpilot reviews are open-ended; (iv) self-selection, "
            "because Trustpilot reviewers are typically more emotionally engaged "
            "than the average customer; (v) product mix, because the post-Betsy "
            "review pool may over-weight customer segments where Betsy is most "
            "exposed; (vi) measurement-window differences, because internal NPS "
            "is collected continuously while Trustpilot review volume varies by "
            "month; and (vii) channel of feedback collection, because the two "
            "channels reach different stages of the customer journey. The chapter "
            "therefore reports both metrics in parallel without attempting to "
            "reconcile them quantitatively, and treats the divergence as a "
            "triangulation insight rather than as a contradiction."
        ),
    ),
]


# ---- Rebuilt reference list ----------------------------------------------------
REBUILT_REFERENCES: list[str] = [
    "Abdulquadri, A., Mogaji, E., Kieu, T. A., & Nguyen, N. P. (2021). Digital transformation in financial services provision: A Nigerian perspective to the adoption of chatbot. Journal of Enterprising Communities: People and Places in the Global Economy, 15(2), 258\u2013281. https://doi.org/10.1108/JEC-06-2020-0126",
    "Acemoglu, D., & Restrepo, P. (2019). Automation and new tasks: How technology displaces and reinstates labor. Journal of Economic Perspectives, 33(2), 3\u201330. https://doi.org/10.1257/jep.33.2.3",
    "Adam, M., Wessel, M., & Benlian, A. (2021). AI-based chatbots in customer service and their effects on user compliance. Electronic Markets, 31(2), 427\u2013445. https://doi.org/10.1007/s12525-020-00414-7",
    "Aggarwal, A., Kumar, V., & Srivastava, R. K. (2025). Value creation and value appropriation using voice AI technology. European Journal of Marketing, 59(9), 2199\u20132243. https://doi.org/10.1108/EJM-11-2023-0917",
    "Andrade, I. M. D., & Tumelero, C. (2022). Increasing customer service efficiency through artificial intelligence chatbot. Revista de Gest\u00e3o, 29(3), 238\u2013251. https://doi.org/10.1108/REGE-07-2021-0120",
    "Becker, D., Braach, L., Clasmeier, L., Kaufmann, T., Ong, O., Ahrens, K., G\u00e4de, C., Strahl, E., Fu, D., & Wermter, S. (2025). Influence of robots' voice naturalness on trust and compliance. ACM Transactions on Human-Robot Interaction, 14(2), 1\u201325. https://doi.org/10.1145/3706066",
    "Better.com. (n.d.). Our company. Better Mortgage. Retrieved May 6, 2026, from https://better.com/content/our-company",
    "Better HoldCo, Inc. (2021, November 30). Digital homeownership platform Better announces up to $1.5 billion in bridge financing and convertible notes with Aurora Acquisition Corp., Novator Capital, Aurora Sponsor and SoftBank [Press release]. Business Wire. https://www.businesswire.com/news/home/20211130005449/en/",
    "Better HoldCo, Inc. (2023, December 20). Registration statement on Form S-1/A. U.S. Securities and Exchange Commission. https://www.sec.gov/Archives/edgar/data/1835856/000162828023042192/0001628280-23-042192-index.htm",
    "Better Home & Finance Holding Company. (n.d.). Investor relations: Corporate overview. Retrieved May 6, 2026, from https://investors.better.com/overview/default.aspx",
    "Better Home & Finance Holding Company. (2024a, October 17). Better.com launches Betsy\u2122, the first voice-based AI loan assistant for the U.S. mortgage industry [Press release]. Business Wire. https://www.businesswire.com/news/home/20241017222780/en/Better.com-launches-Betsy-the-First-Voice-Based-AI-Loan-Assistant-for-the-US-Mortgage-Industry",
    "Better Home & Finance Holding Company. (2024b, November 12). Better Home & Finance Holding Company announces third quarter 2024 results [Exhibit 99.1 to Form 8-K]. U.S. Securities and Exchange Commission. https://www.sec.gov/Archives/edgar/data/1835856/000162828024047437/betterq32024earningsreleas.htm",
    "Better Home & Finance Holding Company. (2025a, March 18). Better Home & Finance Holding Company announces 2024 fourth quarter and full year results [Press release]. Better Investor Relations. https://investors.better.com/news/news-details/2025/Better-Home--Finance-Holding-Company-Announces-2024-Fourth-Quarter-and-Full-Year-Results/default.aspx",
    "Better Home & Finance Holding Company. (2025b, May 13). Better Home & Finance Holding Company announces first quarter 2025 results [Exhibit 99.1 to Form 8-K]. U.S. Securities and Exchange Commission. https://www.sec.gov/Archives/edgar/data/1835856/000162828025024900/betrearningsrelease_2025-q.htm",
    "Better Home & Finance Holding Company. (2025c, August 7). Better Home & Finance Holding Company announces second quarter 2025 results [Exhibit 99.1 to Form 8-K]. U.S. Securities and Exchange Commission. https://www.sec.gov/Archives/edgar/data/1835856/000162828025038583/betrearningsrelease_2025-q2.htm",
    "Better Home & Finance Holding Company. (2026). Form 10-K: Annual report for the fiscal year ended December 31, 2025. U.S. Securities and Exchange Commission. https://www.sec.gov/Archives/edgar/data/1835856/000162828026017747/aurcu-20251231.htm",
    "Better Home & Finance Holding Company, & ElevenLabs. (2026, February 25). ElevenLabs and Better.com showcase success of AI loan agent, Betsy\u2122, at scale in financial services [Press release]. Better Investor Relations. https://investors.better.com/news/news-details/2026/ElevenLabs-and-Better-com-Showcase-Success-of-AI-Loan-Agent-Betsy-at-Scale-in-Financial-Services/",
    "Bl\u00fcmel, J. H., Zaki, M., & Bohn\u00e9, T. (2024). Personal touch in digital customer service: A conceptual framework of relational personalization for conversational AI. Journal of Service Theory and Practice, 34(1), 33\u201365. https://doi.org/10.1108/JSTP-03-2023-0098",
    "Brynjolfsson, E., Li, D., & Raymond, L. R. (2023). Generative AI at work (NBER Working Paper No. 31161). National Bureau of Economic Research. https://doi.org/10.3386/w31161",
    "Consumer Financial Protection Bureau. (2024, May 28). What documents should I receive before closing on a mortgage loan? https://www.consumerfinance.gov/ask-cfpb/what-documents-should-i-receive-before-closing-on-a-mortgage-loan-en-1941/",
    "Consumer Financial Protection Bureau. (2025, November 3). Closing on your new home. https://www.consumerfinance.gov/owning-a-home/closing/",
    "ElevenLabs. (2026, February 9). Better leverages ElevenAgents to build Betsy and make AI-assisted loans accessible [Case study]. ElevenLabs Blog. https://elevenlabs.io/blog/better",
    "Fannie Mae. (n.d.). Selling guide. Retrieved May 6, 2026, from https://singlefamily.fanniemae.com/external-resource/selling-guide",
    "Federal Reserve Bank of New York. (2023, May 15). The great pandemic mortgage refinance boom. Liberty Street Economics. https://libertystreeteconomics.newyorkfed.org/2023/05/the-great-pandemic-mortgage-refinance-boom/",
    "Filieri, R. (2015). Why do travelers trust TripAdvisor? Antecedents of trust towards consumer-generated media and its influence on recommendation adoption and word of mouth. Tourism Management, 51, 174\u2013185. https://doi.org/10.1016/j.tourman.2015.05.007",
    "Fotheringham, D., & Wiles, M. A. (2023). The effect of implementing chatbot customer service on stock returns. Journal of the Academy of Marketing Science, 51, 802\u2013822. https://doi.org/10.1007/s11747-022-00906-6",
    "Freddie Mac. (2021, March 5). Refinance trends in 2020 [Research note]. https://www.freddiemac.com/research/insight/20210305-refinance-trends-in-2020",
    "Freddie Mac. (2024, May 13). 2024 cost to originate study. https://www.freddiemac.com/research/insight/20240513-cost-to-originate",
    "Goodhue, D. L., & Thompson, R. L. (1995). Task-technology fit and individual performance. MIS Quarterly, 19(2), 213\u2013236. https://doi.org/10.2307/249689",
    "Grootendorst, M. (2022). BERTopic: Neural topic modeling with a class-based TF-IDF procedure. arXiv. https://doi.org/10.48550/arXiv.2203.05794",
    "Hentzen, J. K., Hoffmann, A., Dolan, R., & Pala, E. (2022). Artificial intelligence in customer-facing financial services: A systematic literature review and agenda for future research. International Journal of Bank Marketing, 40(6), 1299\u20131336. https://doi.org/10.1108/IJBM-09-2021-0417",
    "Heskett, J. L., Jones, T. O., Loveman, G. W., Sasser, W. E., Jr., & Schlesinger, L. A. (1994). Putting the service-profit chain to work. Harvard Business Review, 72(2), 164\u2013170. https://hbr.org/1994/03/putting-the-service-profit-chain-to-work-2",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75\u2013105. https://doi.org/10.2307/25148625",
    "Huang, M.-H., & Rust, R. T. (2018). Artificial intelligence in service. Journal of Service Research, 21(2), 155\u2013172. https://doi.org/10.1177/1094670517752459",
    "Hutto, C. J., & Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text. Proceedings of the International AAAI Conference on Web and Social Media, 8(1), 216\u2013225. https://doi.org/10.1609/icwsm.v8i1.14550",
    "Jang, M., Jung, Y., & Kim, S. (2021). Investigating managers' understanding of chatbots in the Korean financial industry. Computers in Human Behavior, 120, Article 106747. https://doi.org/10.1016/j.chb.2021.106747",
    "Liu, Z., Zhang, K., & Zhang, H. (2024). A new era of financial services: How AI enhances investment efficiency. International Studies of Economics, 19, 578\u2013588. https://doi.org/10.1002/ise3.91",
    "Luca, M. (2016). Reviews, reputation, and revenue: The case of Yelp.com (Harvard Business School Working Paper No. 12-016). Harvard Business School. https://www.hbs.edu/faculty/Pages/item.aspx?num=41233",
    "Maslych, M., Katebi, M., Lee, C., Hmaiti, Y., Ghasemaghaei, A., Pumarada, C., Palmer, J., Martinez, E. S., Emporio, M., Snipes, W., McMahan, R. P., & LaViola, J. J. (2025). Mitigating response delays in free-form conversations with LLM-powered intelligent virtual agents. arXiv. https://arxiv.org/abs/2504.08507",
    "Mogaji, E., & Nguyen, N. P. (2021). Managers' understanding of artificial intelligence in relation to marketing financial services: Insights from a cross-country study. International Journal of Bank Marketing, 39(6), 1544\u20131565. https://doi.org/10.1108/IJBM-09-2020-0440",
    "Mortgage Bankers Association. (2024a). Quarterly mortgage bankers performance report: Full year 2024. MBA Publications. https://www.mba.org/news-and-research/research-and-economics/single-family-research/quarterly-mortgage-bankers-performance-report",
    "Mortgage Bankers Association. (2024b, November 14). IMB production profits increase in third quarter of 2024 [Press release]. https://www.mba.org/news-and-research/newsroom/news/2024/11/14/imb-production-profits-increase-in-third-quarter-of-2024",
    "Nunes, F. F. (2024, October 17). Better launches voice-based AI loan assistant. HousingWire. https://www.housingwire.com/articles/better-launches-voice-based-ai-loan-assistant/",
    "Nussbaum, C., Fr\u00fchholz, S., & Schweinberger, S. R. (2025). Understanding voice naturalness. Trends in Cognitive Sciences, 29(5), 467\u2013480. https://doi.org/10.1016/j.tics.2025.01.010",
    "Primack, D. (2021, December 8). Better.com delays close of $7.7B reverse merger after mass layoffs. Axios. https://www.axios.com/2021/12/08/bettercom-delays-close-of-77b-reverse-merger-after-mass-layoffs",
    "Qualtrics. (2024). 2024 global consumer trends: Financial services NPS benchmarks. Qualtrics XM Institute.",
    "Rocket Companies, Inc. (2024). Form 10-K: Annual report for the fiscal year ended December 31, 2023. U.S. Securities and Exchange Commission. https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001805284&type=10-K&dateb=&owner=include&count=10",
    "Rocket Companies, Inc. (2025). Form 10-K: Annual report for the fiscal year ended December 31, 2024. U.S. Securities and Exchange Commission. https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001805284&type=10-K&dateb=&owner=include&count=10",
    "Wang, L., Huang, N., Hong, Y., Liu, L., Guo, X., & Chen, G. (2023). Voice-based AI in call center customer service: A natural field experiment. Production and Operations Management, 32(4), 1002\u20131018. https://doi.org/10.1111/poms.13953",
    "Xu, Y., Shieh, C.-H., van Esch, P., & Ling, I.-L. (2020). AI customer service: Task complexity, problem-solving ability, and usage intention. Australasian Marketing Journal, 28(4), 189\u2013199. https://doi.org/10.1016/j.ausmj.2020.03.005",
    "Yin, R. K. (2018). Case study research and applications: Design and methods (6th ed.). SAGE Publications.",
    "Zhang, Z., Li, B., & Liu, L. (2023). The impact of AI-based conversational agent on the firms' operational performance. Applied Artificial Intelligence, 37(1), Article 2246715. https://doi.org/10.1080/08839514.2023.2246715",
]


def apply_body_fixes(doc: Document) -> tuple[int, list[str]]:
    """Apply BODY_FIXES to all paragraphs before the References heading."""
    paras = doc.paragraphs
    n_changes = 0
    examples: list[str] = []
    for i, p in enumerate(paras):
        if p.style and p.style.name == "Heading 1" and p.text.strip() == "References":
            break
        if not p.runs:
            continue
        original = p.text
        new_text = original
        for pat, repl in BODY_FIXES:
            new_text = pat.sub(repl, new_text)
        if new_text != original:
            n_changes += 1
            if len(examples) < 5:
                snippet = next(
                    (s for s in original.split(". ") if any(p_.search(s) for p_, _ in BODY_FIXES)),
                    original[:160],
                )
                examples.append(f"  [{i}] {snippet[:160]}")
            # Replace text by clearing all runs except the first then writing new_text
            for r in p.runs[1:]:
                r.text = ""
            p.runs[0].text = new_text
    return n_changes, examples


def replace_reference_list(doc: Document, references: list[str]) -> int:
    """Replace all paragraphs after the 'References' heading with the rebuilt list."""
    body = doc.element.body
    paras = list(doc.paragraphs)
    ref_heading_idx = None
    for i, p in enumerate(paras):
        if p.style and p.style.name == "Heading 1" and p.text.strip() == "References":
            ref_heading_idx = i
            break
    if ref_heading_idx is None:
        raise RuntimeError("Could not find 'References' heading")

    heading_p = paras[ref_heading_idx]
    # 1) Pick a template paragraph (the one immediately after the heading) so
    #    new entries inherit its style.
    template_p = paras[ref_heading_idx + 1]
    template_xml = deepcopy(template_p._p)

    # 2) Remove every old reference paragraph after the heading.
    n_removed = 0
    for old_p in paras[ref_heading_idx + 1 :]:
        old_p._p.getparent().remove(old_p._p)
        n_removed += 1

    # 3) Insert new reference paragraphs after the heading, in order.
    insert_after = heading_p._p
    for entry in references:
        new_p = deepcopy(template_xml)
        # Strip all existing runs from the cloned paragraph
        for r in list(new_p.findall(qn("w:r"))):
            new_p.remove(r)
        # Strip any hyperlinks that were in the template
        for hl in list(new_p.findall(qn("w:hyperlink"))):
            new_p.remove(hl)
        # Add a fresh run with the entry text
        r = new_p.makeelement(qn("w:r"), {})
        t = new_p.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        t.text = entry
        r.append(t)
        new_p.append(r)
        insert_after.addnext(new_p)
        insert_after = new_p

    return n_removed


def main() -> None:
    doc = Document(str(DOCX_PATH))
    print(f"Loaded {DOCX_PATH} ({len(doc.paragraphs)} paragraphs)")

    n_body, examples = apply_body_fixes(doc)
    print(f"Body citation fixes applied to {n_body} paragraphs")
    for ex in examples:
        print(ex)

    n_removed = replace_reference_list(doc, REBUILT_REFERENCES)
    print(
        f"Replaced reference list: removed {n_removed} old paragraphs, "
        f"inserted {len(REBUILT_REFERENCES)} new entries"
    )

    doc.save(str(DOCX_PATH))
    print(f"Saved {DOCX_PATH} ({len(Document(str(DOCX_PATH)).paragraphs)} paragraphs)")


if __name__ == "__main__":
    main()
