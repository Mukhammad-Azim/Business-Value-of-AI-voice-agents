"""
apply_thesis_cleanup.py
=======================

Pre-Chapter-5 quality-control and consistency cleanup pipeline for the MSc thesis
"Business Value of Voice LLM Agents in Financial Services: Evidence from Better.com."

This script reads `04_Drafts/Main_Draft.docx`, applies a deterministic set of
edits derived from the supervisor-style cleanup brief, and writes the result
back to the same path. Run from the repo root:

    python3 08_Code/apply_thesis_cleanup.py

The edits implemented here cover (numbering matches the cleanup brief):

  2.  RQ2 wording standardization
  3.  Replacement of "DiD-style" / "Difference-in-Differences" with
      descriptive benchmark language
  4.  Financial-methodology wording (10 SEC filings, 13x15 dataset, manual
      transcription, no OCR, etc.)
  5.  Finalized Better financial values
  6.  Total_Expenses clarification + cross-company comparison caveat
  7.  Operational Leverage Index restricted to appendix; value fixed to 195.3
  8.  NLP dataset consistency (1,934 raw / 1,915 clean / 1,746 pre /
      169 post / topic-modeling 1,804)
  9.  NLP sentiment + rating values (full precision, Cohen's d included)
  10. NLP keyword prevalence values and Bonferroni framing
  11. AI-invisibility wording (no "proven finding" framing)
  12. NPS interpretation (management-reported, medium reliability, divergence)
  13. BERTopic interpretation (descriptive only, n=121 caveat)
  14. Chapter 4 cleanup (descriptive, no Chapter 5 findings)
  15. Chapter 3 methodology restructure into 3.1-3.8 sections, with an
      operationalization-of-business-value table
  16. Chapter 1 academic-tone softening
  18. Figure-placement notes alongside relevant paragraphs

The final pass also enforces a consistent causality-language register
("consistent with", "associated with", "directionally supports", etc.).

The edits are conservative: they replace existing paragraph text but preserve
paragraph styles where possible, and only insert new paragraphs at clearly
defined anchor points so that document structure remains intact.
"""

from __future__ import annotations

import os
import sys
from copy import deepcopy

import docx
from docx.oxml.ns import qn
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace the entire text of a paragraph, keeping its style.

    python-docx exposes a writable `text` attribute. Setting it clears all
    runs and creates one new run with the given text. This loses any
    inline formatting inside the paragraph but preserves the paragraph
    style (heading level, list level, etc.). For pure-prose paragraphs in
    this draft that is acceptable.
    """
    paragraph.text = new_text


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    """Replace `old` with `new` inside a paragraph, preserving formatting
    where possible.

    Returns True if any replacement was performed.
    """
    # Try run-by-run replacement first
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    if changed:
        return True
    # Fall back to whole-paragraph replacement if the old string spans
    # multiple runs (which loses run-level formatting but keeps paragraph
    # style).
    full = paragraph.text
    if old in full:
        paragraph.text = full.replace(old, new)
        return True
    return False


def insert_paragraph_after(paragraph, text: str, style_name: str | None = None):
    """Insert a new paragraph with the given text directly after the
    supplied paragraph and return the new paragraph object."""
    new_p = paragraph._parent.add_paragraph(text, style=style_name) if False else None
    # The above shortcut appends to end; we need positional insertion.
    new_p_xml = docx.oxml.OxmlElement("w:p")
    paragraph._p.addnext(new_p_xml)
    new_para = docx.text.paragraph.Paragraph(new_p_xml, paragraph._parent)
    if style_name:
        new_para.style = paragraph.part.document.styles[style_name]
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, headers: list[str], rows: list[list[str]]):
    """Insert a simple table directly after `paragraph` and return the
    table object."""
    doc = paragraph.part.document
    # Build the table at the end of the document (python-docx default)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        # Some docx files do not register the Table Grid style. Fall
        # back to default.
        pass
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    # Body rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val
    # Move the table from end-of-document to immediately after `paragraph`
    table_xml = table._tbl
    table_xml.getparent().remove(table_xml)
    paragraph._p.addnext(table_xml)
    return table


# ---------------------------------------------------------------------------
# Edit pipeline
# ---------------------------------------------------------------------------

# Final RQ2 wording (verbatim from the supervisor brief)
FINAL_RQ2 = (
    "RQ2: What measurable business value can be associated with Better.com\u2019s "
    "voice AI deployment across operational efficiency, productivity, customer "
    "experience, and financial performance, and what factors mediate the "
    "realization of this value?"
)


# Whole-paragraph rewrites keyed by the leading prefix of the existing
# paragraph text. Using a prefix match (rather than full-text equality)
# tolerates minor whitespace differences. The order matches the order
# paragraphs appear in the document.
PARAGRAPH_REWRITES: list[tuple[str, str]] = [
    # ------------------------------------------------------------
    # Chapter 1 - tone softening (Task 16)
    # ------------------------------------------------------------
    (
        "While the impact of artificial intelligence is ubiquitous,",
        "While the impact of artificial intelligence is broad, the financial "
        "services industry is a particularly informative empirical setting for "
        "these technological advancements. The financial sector is "
        "characterized by high-volume, high-stakes customer interactions. "
        "Whether facilitating mortgage origination, managing retail banking "
        "queries, or processing insurance claims, financial institutions "
        "operate under sustained pressure to deliver consistent customer "
        "experiences while managing operational costs and margin compression. "
        "The sector is also heavily regulated, requiring customer "
        "interactions to maintain standards of accuracy, transparency, and "
        "auditability (Hentzen et al., 2022).",
    ),
    (
        "More recently, financial organizations have begun deploying AI directly",
        "More recently, financial organizations have begun deploying AI "
        "directly at the customer interface, using automated routing and "
        "rudimentary text-based chatbots to manage high query volumes and "
        "provide continuous service availability (Mogaji & Nguyen, 2021). "
        "Yet AI's organizational impact extends beyond simple cost reduction. "
        "As Liu et al. (2024) observe, AI integration affects efficiency, "
        "risk management, and the overall quality of customer interactions. "
        "In financial services, where consumer trust and clarity are central, "
        "the limitations of legacy digital interfaces have become increasingly "
        "apparent. Complex financial decisions, such as securing a 30-year "
        "residential mortgage, frequently require reassurance, empathetic "
        "communication, and dynamic problem-solving that standard, rigid "
        "automated systems are not well suited to provide.",
    ),
    (
        "Within this rapidly evolving landscape, the introduction of voice-based",
        "Within this evolving landscape, the introduction of voice-based "
        "Large Language Model (LLM) agents represents a significant "
        "technological shift relative to preceding customer service "
        "technologies. Voice LLM agents are conversational systems that "
        "integrate automatic speech recognition (ASR), generative natural "
        "language processing (NLP), and near-instantaneous text-to-speech "
        "(TTS) synthesis. They differ meaningfully from first-generation "
        "digital assistants by operating as context-aware, adaptive digital "
        "agents capable of navigating multi-turn human dialogues with natural "
        "prosody and inflection.",
    ),
    (
        "From an academic perspective, this study directly addresses the",
        "From an academic perspective, this study directly addresses an "
        "important empirical gap surrounding the actual organizational "
        "outcomes of advanced conversational AI. By systematically connecting "
        "technical AI implementation to measurable business performance "
        "outcomes, the research extends the theoretical understanding of AI "
        "value creation. It moves the academic conversation beyond the "
        "prevalent focus on user acceptance and algorithmic design and "
        "establishes an empirical link between generative voice AI and the "
        "underlying economics of customer service. By utilizing a real-world, "
        "high-stakes financial services case, it grounds theoretical "
        "discussions of AI transformation in observable corporate realities, "
        "contributing empirical data to the literature on digital business "
        "strategy and information systems management (Liu et al., 2024; "
        "Aggarwal et al., 2025).",
    ),
    # ------------------------------------------------------------
    # Chapter 1 - RQ2 wording (Task 2)
    # ------------------------------------------------------------
    (
        "RQ2: What measurable business value",
        FINAL_RQ2,
    ),
    # ------------------------------------------------------------
    # Chapter 1 - methodological-contribution paragraph
    # (drop "unique and valuable", "ensures a validity")
    # ------------------------------------------------------------
    (
        "Perhaps most notably, this thesis delivers a highly robust",
        "Methodologically, this thesis contributes a mixed-methods "
        "triangulation design that combines qualitative process tracing of "
        "the Better.com case study with longitudinal financial analysis of "
        "SEC-reported performance metrics and quantitative Natural Language "
        "Processing (NLP) analysis of customer review data. The combination "
        "of regulated corporate financial disclosures with algorithmic "
        "sentiment and topic modeling of unsolicited consumer feedback is "
        "empirically valuable for contemporary information systems research "
        "and allows the study to examine whether operational improvements "
        "are accompanied by changes in independent customer-experience "
        "signals.",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.2 Research Design (Task 15)
    # P109 paragraph: replace "test the hypothesis" with "examine"
    # ------------------------------------------------------------
    (
        "To operationalize this approach, the study employs a Convergent",
        "To operationalize this approach, the study employs a Convergent "
        "Parallel Mixed Methods Case Study design, guided by Yin (2018). In "
        "a convergent parallel design, quantitative and qualitative data "
        "streams are collected and analyzed independently, and results are "
        "merged only during the final interpretation phase to cross-validate "
        "and triangulate findings. The design comprises three integrated "
        "strands: (1) a Qualitative Case Study Strand that process-traces "
        "the historical, organizational, and technical integration of the "
        "Tinman platform and Betsy voice agent through corporate "
        "communications and technical documentation, primarily addressing "
        "RQ1; (2) a Quantitative Financial Strand that uses longitudinal "
        "SEC filings to examine evidence consistent with AI-enabled "
        "operational leverage, addressing RQ2; and (3) a Computational NLP "
        "Strand that applies sentiment analysis, keyword prevalence, and "
        "BERTopic topic modeling to unsolicited Trustpilot reviews to capture "
        "the customer-experience dimension of RQ2. Rocket Mortgage is used "
        "throughout as a directional industry benchmark rather than as a "
        "formal counterfactual.",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.4 Data Sources / Financial filings (Task 4)
    # P116
    # ------------------------------------------------------------
    (
        "The financial analysis is anchored in thirteen annual observations",
        "The financial analysis is anchored in thirteen annual observations "
        "extracted from ten official SEC regulatory filing PDFs accessed via "
        "the EDGAR public database. The Better.com corpus comprises three "
        "annual 10-K reports (2023, 2024, 2025) and one pre-merger S-1 "
        "Registration Statement. Because Better.com qualifies as a Smaller "
        "Reporting Company, its 10-K filings provide only two years of "
        "audited P&L data; the S-1 was therefore used to bridge the 2021 "
        "data gap and establish the pre-restructuring baseline. The Rocket "
        "Companies benchmark corpus comprises six 10-K filings (2020, 2021, "
        "2022, 2023, 2024, 2025), with Rocket's 2018-2019 values taken from "
        "historical comparative tables in Rocket's 2020 10-K rather than "
        "from separate locally stored PDFs.",
    ),
    # P117 - extraction wording
    (
        "Data extraction employed the Python library pdfplumber",
        "Text extraction and document navigation used the Python libraries "
        "pdfplumber (script: extract_pdf_texts.py) and pypdf (script: "
        "extract_s1.py for the structurally distinct S-1 filing). No optical "
        "character recognition was used, and the pipeline did not "
        "automatically parse financial tables into final numeric data. "
        "Because pdfplumber's extract_text() preserves reading order without "
        "reconstructing tabular structure, all final financial values were "
        "manually identified, standardized to millions of USD, and "
        "cross-verified against the source PDFs. Derived metrics (Expense "
        "Ratio, Revenue per Employee, Profit Margin, Equity Ratio, Revenue "
        "Growth) were then calculated deterministically in Python. The "
        "resulting longitudinal dataset (financial_dataset_annual.csv) "
        "contains 13 rows and 15 columns; financial values are in USD "
        "millions and employee count is reported as headcount. Provenance "
        "for each extracted figure, including the reliance on Rocket's 2020 "
        "10-K for 2018-2019 historical comparatives, is documented in the "
        "accompanying financial_dataset_documentation.md file.",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.4 Trustpilot (Task 8)
    # P119
    # ------------------------------------------------------------
    (
        "To capture the customer perception dimension, an extensive corpus",
        "To capture the customer perception dimension, an extensive corpus "
        "of unsolicited customer reviews was collected from Better.com's "
        "official Trustpilot profile. Trustpilot is widely used in academic "
        "IS research as a proxy for electronic word-of-mouth and service "
        "quality perception (Filieri, 2015; Luca, 2016). Data extraction "
        "was executed using the Apify cloud-based web scraping infrastructure "
        "(Actor ID: l3wcDhSSC96LBRUpc), collecting the unique review ID, "
        "ISO 8601 publication timestamp, full text body, star rating, and "
        "language metadata. The automated scraping yielded a raw corpus of "
        "1,934 reviews spanning August 22, 2020 to April 25, 2026.",
    ),
    # P120 - cleaning + 1,915 + 1,804 (not 1,805)
    (
        "A reproducible data cleaning pipeline was then executed in Python",
        "A reproducible data cleaning pipeline was then executed in Python "
        "(cleaning_reviews.ipynb), applying sequential filters: removal of "
        "18 short reviews (under 20 characters), retention of "
        "English-language reviews only (removing 1 non-English entry), and "
        "removal of exact duplicates. The result was a final clean dataset "
        "of 1,915 reviews. The full clean dataset is the input for sentiment "
        "analysis, star-rating comparisons, and keyword prevalence. For "
        "BERTopic topic modeling only, an additional CEO/scandal-related "
        "outlier filter was applied to remove reviews generated by the 2021 "
        "mass-layoff media controversy that focused on the CEO rather than "
        "the mortgage product, producing a topic-modeling input of 1,804 "
        "documents. The CEO/scandal filter is not applied to the rating, "
        "sentiment, or keyword analyses, which use the full 1,915-review "
        "dataset. Pre-Betsy and Post-Betsy periods are defined relative to "
        "Betsy's October 17, 2024 launch date (Pre: 1,746 reviews; Post: "
        "169 reviews).",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.5 Financial methodology (Task 3, 4, 6, 7)
    # P125 (operational leverage hypothesis -> examine evidence consistent with)
    # ------------------------------------------------------------
    (
        "The financial analysis employed a longitudinal framework engineered",
        "The financial analysis used a longitudinal framework designed to "
        "examine evidence consistent with AI-enabled operational leverage: "
        "the proposition that AI deployment is associated with revenue "
        "generation that grows faster than the operational expense base, "
        "without proportional human capital expansion. Variables were "
        "operationalized and calculated via a Python pipeline. The Expense "
        "Ratio (Total Expenses / Total Net Revenue) is used as a broad "
        "cost-intensity indicator. Total Expenses is treated as a "
        "standardized total expense base constructed from reported financial "
        "statements, rather than as a strict line-item operating-expense "
        "measure. Because Better.com and Rocket Companies use different "
        "accounting presentations, cross-company expense-ratio comparisons "
        "should be interpreted directionally rather than as precise "
        "like-for-like measurements. Revenue per Employee (Total Net Revenue "
        "/ Total Headcount) is used as the primary proxy for labor "
        "productivity and AI scalability. Net loss trajectory and profit "
        "margin are taken directly from the Statements of Operations to "
        "track holistic financial performance.",
    ),
    # P126 (Difference-in-Differences -> descriptive benchmark)
    (
        "To isolate Better.com\u2019s performance from the broader macroeconomic",
        "To contextualize Better.com's performance against the broader "
        "macroeconomic environment, the analysis implements a descriptive "
        "benchmark comparison against Rocket Mortgage, contrasting the "
        "pre-to-post change in Better.com's key metrics with the concurrent "
        "change in Rocket's. This comparison is descriptive rather than a "
        "formal econometric causal identification strategy. Better and "
        "Rocket differ substantially in scale (approximately 40 times "
        "revenue), business model, and 2025 acquisition context, so Rocket "
        "is treated as a directional industry benchmark rather than as a "
        "formal counterfactual. Rocket's 2025 data reflects the Mr. Cooper "
        "acquisition, which distorts per-employee comparisons. The "
        "benchmark-adjusted descriptive comparison therefore provides "
        "directional rather than controlled causal evidence. Three "
        "alternative baseline checks are conducted (2024-2025 as the "
        "primary post-transition window, 2023-2025 as supporting "
        "recovery-from-trough evidence, and pre-period-average against "
        "2025 as a robustness sensitivity check) to ensure that findings "
        "are not artifacts of trough-year baseline selection.",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.6 NLP (Task 8, 10, 11)
    # P128 - intro
    # ------------------------------------------------------------
    (
        "The computational analysis of the 1,915 Trustpilot reviews",
        "The computational analysis of the 1,915 Trustpilot reviews uses "
        "three complementary techniques: lexicon-based VADER sentiment "
        "analysis on the full clean dataset, keyword prevalence analysis on "
        "the full clean dataset, and BERTopic neural topic modeling on the "
        "1,804-document CEO/scandal-filtered subset. The combination of "
        "sentiment, keyword, and topic-level views is designed to surface "
        "different facets of customer experience rather than to assert any "
        "single causal claim about Betsy. All analyses are executed in "
        "Python (better_reviews_topic_modeling.ipynb).",
    ),
    # P130 - BERTopic
    (
        "Neural Topic Modeling (BERTopic).",
        "Neural Topic Modeling (BERTopic). To discover latent thematic "
        "structure without imposing a priori researcher categories, the "
        "study used BERTopic (Grootendorst, 2022), which combines transformer "
        "embeddings with density-based clustering. The pipeline proceeded in "
        "four steps: (1) the all-MiniLM-L6-v2 Sentence-Transformer model "
        "converted each review into a dense semantic vector; (2) UMAP "
        "(n_neighbors=10, n_components=5, random_state=42) compressed "
        "embeddings into a lower-dimensional space; (3) HDBSCAN "
        "(min_cluster_size=10, min_samples=3) identified dense document "
        "regions while sequestering outliers; and (4) class-based TF-IDF "
        "extracted representative keywords per cluster, enabling researcher-"
        "assigned human-readable labels. The model produced 44 topics from "
        "1,804 input documents, with 1,198 documents successfully clustered "
        "and 606 assigned to the HDBSCAN noise class. The clustered "
        "non-outlier subsample used for topic-share comparisons therefore "
        "consists of 1,077 pre-Betsy and 121 post-Betsy reviews. Because "
        "the post-Betsy clustered non-outlier sample is small, topic-share "
        "changes are interpreted as descriptive signals rather than as "
        "statistically tested population-level changes.",
    ),
    # P131 - keyword prevalence
    (
        "Keyword Prevalence and AI Invisibility Testing.",
        "Keyword Prevalence and Explicit AI Mention Analysis. Beyond topic "
        "discovery, a targeted keyword prevalence analysis measured the "
        "salience of seven thematic groups: Explicit AI/Automation, "
        "Speed/Efficiency, Human Support, Communication, Delay/Friction, "
        "Cost/Rates, and Digital/Process. Using strict word-boundary regular "
        "expressions, reviews were binary-coded for keyword presence, and "
        "Pearson's chi-squared tests evaluated whether prevalence shifted "
        "between the pre-Betsy and post-Betsy periods. A Bonferroni "
        "correction was applied to mitigate Type I error risk from multiple "
        "tests. Under the unadjusted alpha, Cost/Rates (Pre 28.92%, Post "
        "16.57%, change -12.35 pp, p = 0.0009) and Digital/Process (Pre "
        "32.65%, Post 23.67%, change -8.98 pp, p = 0.0211) shift "
        "significantly; only Cost/Rates survives strict Bonferroni "
        "correction. Speed/Efficiency (Pre 52.00%, Post 44.38%, change "
        "-7.63 pp, p = 0.0698) is suggestive but not significant. Human "
        "Support (Pre 43.81%, Post 40.83%), Communication (Pre 32.53%, "
        "Post 30.77%), and Delay/Friction (Pre 13.52%, Post 10.06%) do not "
        "differ significantly. Explicit AI/Automation mentions remain near "
        "zero in both periods (Pre 1.15%, Post 1.18%; Fisher's exact p = "
        "1.0000). The near-zero explicit AI mention rate suggests that "
        "customers do not usually frame their experience in AI-specific "
        "terms, supporting the interpretation of voice AI as invisible "
        "infrastructure. Because the keyword analysis cannot determine "
        "whether individual customers actually interacted with Betsy, this "
        "is a descriptive observation rather than a measure of customer "
        "exposure to the AI agent.",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.7 Operationalization (Task 15)
    # P135 (rewrite + add table afterwards)
    # ------------------------------------------------------------
    (
        "A critical methodological component is the operationalization",
        "A central methodological step is the operationalization of "
        "business value into measurable, contextually accurate indicators. "
        "Rather than treating value creation as an abstract construct, the "
        "research relies on indicators specific to digital mortgage lending "
        "and grounded in standardized SEC accounting definitions to minimize "
        "subjective researcher bias. The mapping between business-value "
        "dimensions and indicators is summarized in the operationalization "
        "table below. Operational efficiency is captured by the Expense "
        "Ratio, the Cost-to-Originate estimate, and expense growth. "
        "Productivity and scalability are captured by Revenue per Employee, "
        "headcount versus revenue trends, and Betsy's reported call volume. "
        "Financial performance is captured by revenue, net loss, and profit "
        "margin. Customer experience is captured by management-reported "
        "NPS, Trustpilot star ratings, VADER sentiment scores, BERTopic "
        "topic shares, and keyword prevalence. Integration mechanisms are "
        "captured qualitatively through the Tinman platform, Betsy as "
        "voice layer, ElevenLabs as voice infrastructure provider, and the "
        "human escalation model. Management-reported NPS (39\u219264) is treated "
        "as a medium-reliability indicator and is interpreted alongside "
        "independent Trustpilot evidence rather than as definitive proof "
        "that customers preferred Betsy. The author-constructed Operational "
        "Leverage Index (OLI; 2023=100, 2025=195.3) is retained as a "
        "supplementary, equally-weighted composite visualization that is "
        "sensitive to baseline and weighting assumptions; it is reported in "
        "the appendix and is not used as primary evidence in Chapter 5.",
    ),
    # ------------------------------------------------------------
    # Chapter 3 - 3.7/3.8 Triangulation Logic (Task 9, 11, 12, 13)
    # P138 - update with finalized values + invisibility wording
    # ------------------------------------------------------------
    (
        "Crucially, triangulation is designed to handle divergence",
        "Triangulation is designed to handle divergence rather than suppress "
        "it. The thesis exhibits a meaningful divergence across measurement "
        "instruments. Management reported an internally measured NPS "
        "improvement from 39 to 64 following Betsy's launch, and audited "
        "financial data shows substantial expense-intensity improvement and "
        "revenue recovery. In parallel, the independent NLP analysis shows "
        "a statistically significant decline in Trustpilot star ratings "
        "(pre-Betsy mean 4.2806 vs post-Betsy mean 3.8580, difference "
        "-0.4226, Welch's t-test p = 0.0019, Cohen's d = 0.2899) and in "
        "VADER compound sentiment (pre-Betsy mean 0.6416 vs post-Betsy "
        "mean 0.5151, difference -0.1265, Welch's t-test p = 0.0057, "
        "Cohen's d = 0.2505) over the same period. Rather than discarding "
        "this conflicting signal, the mixed-methods architecture treats it "
        "as an analytically valuable triangulation insight: management-"
        "reported NPS is medium-reliability evidence, while Trustpilot "
        "discourse is independent and self-selected, so divergence is "
        "consistent with different measurement instruments capturing "
        "different customer-experience dimensions. BERTopic topic shifts "
        "(notably HELOC complaints +4.22 pp, Smooth Transaction +4.03 pp, "
        "Responsive Service +3.93 pp, Closing Timeline +3.55 pp, refinance "
        "topics combined -7.66 pp, Home Purchase -3.48 pp) provide a "
        "descriptive product-mix interpretation but, because the post-Betsy "
        "clustered non-outlier sample is only 121 reviews, are not "
        "interpreted as statistically tested population-level changes. "
        "This triangulation logic prevents both uncritical acceptance of "
        "corporate marketing narratives and over-interpretation of any "
        "single metric.",
    ),
    # P144 - NLP imbalance paragraph (already mostly OK, but tidy)
    (
        "NLP dataset imbalance and selection bias.",
        "NLP dataset imbalance and selection bias. Analyzing an "
        "approximately 18-month post-Betsy window against a 4-year "
        "pre-Betsy history produces an imbalanced sample (Pre n = 1,746; "
        "Post n = 169, ratio 10.3:1). While Welch's t-test and the "
        "chi-squared test technically accommodate unequal group sizes, the "
        "small post-period sample carries wider confidence intervals and is "
        "sensitive to idiosyncratic reviews. The BERTopic clustered "
        "non-outlier sub-sample is smaller still (Pre n = 1,077; Post n = "
        "121), so topic-share changes are descriptive signals rather than "
        "tested population-level changes. Trustpilot reviewers are also "
        "self-selected and prone to extreme-response bias, meaning the "
        "corpus may over-represent highly satisfied or highly frustrated "
        "customers relative to the median user experience.",
    ),
    # ------------------------------------------------------------
    # Chapter 4 - financial values consistency (Task 5, 14)
    # P234 (2023 revenue $77M -> $76.8M)
    # ------------------------------------------------------------
    (
        "The broader mortgage market also changed materially during this",
        "The broader mortgage market also changed materially during this "
        "period. After the 2020-2021 refinancing boom, rising interest rates "
        "reduced refinancing demand and made origination volumes more "
        "difficult to sustain. Better's 2023 annual results explicitly "
        "describe 2023 as one of the most challenging mortgage macro "
        "environments in recent history, with average 30-year fixed mortgage "
        "rates around 7%. In the same release, Better reported 2023 revenue "
        "of approximately $76.8 million, funded loan volume of $3.0 billion, "
        "and an adjusted EBITDA loss of $163 million, while emphasizing "
        "that it had reduced total expenses by 71% year over year and by "
        "more than $1.1 billion compared with 2021 (Better Home & Finance "
        "Holding Company, 2024b).",
    ),
    # P239 (2024 results paragraph - $108M / $72M / $206M / $536M -> finalized)
    (
        "Better\u2019s 2024 financial results reinforce this interpretation.",
        "Better's 2024 financial results reinforce this interpretation. "
        "On a standardized analytical basis, full-year 2024 revenue was "
        "approximately $108.5 million (up from approximately $76.8 million "
        "in 2023) and the net loss narrowed to approximately $206.3 million "
        "(from approximately $536.4 million in 2023). Funded loan volume "
        "increased to approximately $3.6 billion from $3.0 billion in 2023, "
        "while total loans increased to approximately 11,800 from "
        "approximately 8,600 (Better Home & Finance Holding Company, "
        "2025a). These improvements provide descriptive evidence of "
        "Better's recovery trajectory but cannot isolate Betsy's specific "
        "contribution because most of 2024 preceded the launch and "
        "reflects the broader Betsy/Tinman AI operating model and "
        "concurrent restructuring rather than a clean post-Betsy window.",
    ),
]


# Targeted in-text substitutions applied across every paragraph.
SUBSTITUTIONS: list[tuple[str, str]] = [
    # Causality language register (Task 1, 11, 12, 13)
    ("validates the hypothesis", "is consistent with the hypothesis"),
    ("definitively demonstrates", "provides descriptive evidence consistent with"),
    ("transformative results achieved through AI",
     "operational changes consistent with AI-enabled operational leverage"),
    ("only achievable through AI substitution",
     "consistent with AI-enabled operational leverage"),
    ("positive causal interpretation",
     "associative interpretation consistent with AI-enabled operational leverage"),
    # Generic dramatic-language softening (Task 16)
    ("profound and irreversible", "significant"),
    ("massive blind spot", "important gap"),
    ("desperately short", "limited"),
    ("exceptionally rare and incredibly powerful", "empirically valuable"),
    ("virtually infinite scalability", "scalable"),
    ("glaringly apparent", "increasingly apparent"),
    ("glaring empirical void", "important empirical gap"),
    ("exceptionally high volumes of complex",
     "high volumes of complex"),
    ("exceptionally high. The entire data processing pipeline",
     "high. The entire data processing pipeline"),
    ("Computational reliability and reproducibility are exceptionally high",
     "Computational reliability and reproducibility are high"),
    # Causality language additions (Task 1, 11)
    ("Betsy\u2019s contribution is experienced as faster responses, "
     "smoother transactions, and better availability",
     "the near-zero explicit AI mention rate suggests that customers "
     "do not usually frame their experience in AI-specific terms"),
    ("test the operational leverage hypothesis",
     "examine evidence consistent with AI-enabled operational leverage"),
    ("test the hypothesis of AI-driven operational leverage",
     "examine evidence consistent with AI-enabled operational leverage"),
    # AI Invisibility hypothesis -> AI Invisibility interpretation
    ("AI Invisibility hypothesis", "AI invisibility interpretation"),
    # Replace "this thesis is philosophically empowered" softer
    ("philosophically empowered to utilize",
     "well-suited to utilize"),
    # DiD residue (Task 3)
    ("Difference-in-Differences (DiD) style descriptive comparison",
     "descriptive benchmark comparison against Rocket Mortgage"),
    ("Difference-in-Differences (DiD)", "descriptive benchmark comparison"),
    ("Difference-in-Differences", "descriptive benchmark comparison"),
    ("DiD-style", "descriptive benchmark"),
    (" DiD ", " descriptive benchmark comparison "),
    # ------------------------------------------------------------
    # Second-pass corrections (supervisor feedback)
    # ------------------------------------------------------------
    # Repeated phrase left over from the first pass
    ("descriptive benchmark descriptive benchmark comparison",
     "descriptive benchmark comparison"),
    ("descriptive benchmark descriptive benchmark",
     "descriptive benchmark"),
    # "controls for broad market recovery" -> "contextualizes against ..."
    ("comparison with Rocket Mortgage controls for broad market recovery effects, "
     "preventing false association of macro-recovery to the AI agent",
     "comparison with Rocket Mortgage contextualizes Better.com's trajectory "
     "against broad market recovery effects, reducing the risk of attributing "
     "macroeconomic recovery to the voice AI agent"),
    ("controls for broad market recovery effects",
     "contextualizes Better.com's trajectory against broad market recovery effects"),
    ("controls for broad market recovery",
     "contextualizes against broad market recovery"),
    # GDPR softening
    ("ensuring zero risk of GDPR violation or data privacy misuse",
     "minimizing GDPR exposure because no personal data was collected, "
     "linked to identifiable individuals, or stored locally beyond the public "
     "review text"),
    ("zero risk of GDPR violation or data privacy misuse",
     "minimal GDPR exposure because no personally identifiable data was collected"),
    ("zero risk of GDPR violation",
     "minimal GDPR exposure"),
    # Typo fixes ("im" prefix that crept in during a prior copy-edit)
    ("imdemonstrating", "demonstrating"),
    ("imindicates", "indicates"),
    ("imindicate", "indicate"),
    ("imdemonstrate", "demonstrate"),
    # Chapter 1 - additional tone softening (second pass)
    ("a significant transformation across global industries, catalyzing a "
     "paradigm shift",
     "an important transformation across global industries, contributing "
     "to a shift"),
    ("represents a discontinuous technological leap",
     "represents a substantial technological advance"),
    ("decisively shifted from back-office data processing to front-line "
     "customer engagement, fundamentally redefining how organizations interact "
     "with their client bases",
     "shifted from back-office data processing toward front-line customer "
     "engagement, reshaping how organizations interact with their client bases"),
    ("In this unforgiving environment",
     "In this competitive environment"),
    ("a core driver of competitive advantage, customer retention, and "
     "long-term revenue generation",
     "an important driver of competitive position, customer retention, and "
     "revenue generation"),
    ("acutely aware that friction in communication directly correlates with",
     "aware that friction in communication is associated with"),
    ("the industry has historically been an early and aggressive adopter of "
     "AI technologies",
     "the industry has historically been an early adopter of AI technologies"),
    ("Traditional machine learning architectures have already been extensively "
     "deployed for backend operational functions, transforming areas such as",
     "Traditional machine learning architectures have already been deployed "
     "for backend operational functions, supporting areas such as"),
    ("the ideal, high-stakes context in which to evaluate how sophisticated "
     "AI systems transition from technical novelties into core drivers of "
     "measurable business value",
     "an empirically informative, high-stakes context in which to evaluate "
     "how voice-based AI systems are associated with measurable business value"),
    ("represents a significant technological shift relative to preceding "
     "customer service technologies",
     "represents an important step relative to preceding customer service "
     "technologies"),
    ("a critical and widening void remains in the academic literature",
     "an important and growing gap remains in the academic literature"),
    ("the overwhelming majority of contemporary AI research is structurally "
     "skewed toward",
     "much of the contemporary AI research has concentrated on"),
    ("What is fundamentally missing is empirical, business-focused research",
     "What remains underexplored is empirical, business-focused research"),
    ("The existing literature is replete with theoretical assertions",
     "The existing literature contains theoretical assertions"),
    ("to move the academic discourse from theoretical potential to empirical "
     "reality",
     "to move the academic discourse from theoretical claims toward empirical "
     "evidence"),
    ("severe macroeconomic headwinds and extreme margin pressures",
     "significant macroeconomic headwinds and margin pressures"),
    ("highly scalable, AI-driven automation framework",
     "scalable, AI-supported workflow"),
    ("a unique and invaluable opportunity to evaluate real-world business "
     "impact",
     "an empirically informative opportunity to evaluate real-world business "
     "outcomes"),
    ("This research gap represents a significant and pressing barrier",
     "This research gap is an important barrier"),
    ("a research gap in our comprehension of digital transformation",
     "a gap in current understanding of digital transformation"),
    # Capitalization / heading text
    ("Literature review", "Literature Review"),
]


# ---------------------------------------------------------------------------
# Pipeline driver
# ---------------------------------------------------------------------------

def _normalize_ws(s: str) -> str:
    """Normalize whitespace so that NBSPs and other Unicode spaces don't
    defeat prefix matching."""
    # Treat NBSP, narrow-NBSP, en-space, em-space, etc. as a regular space.
    table = {
        ord("\u00a0"): " ",
        ord("\u202f"): " ",
        ord("\u2007"): " ",
        ord("\u2008"): " ",
        ord("\u2009"): " ",
        ord("\u200a"): " ",
        ord("\u2028"): " ",
        ord("\u2029"): " ",
    }
    return s.translate(table)


def find_paragraph_by_prefix(doc, prefix: str):
    norm_prefix = _normalize_ws(prefix)
    for para in doc.paragraphs:
        if _normalize_ws(para.text).startswith(norm_prefix):
            return para
    return None


def normalize_chapter_headings(doc) -> None:
    """Apply consistent chapter numbering and heading levels.

    - Heading 1 chapter titles get sequential numeric prefixes
      (1. Introduction, 2. Literature Review, 3. Research
      Methodology, 4. Case Study Description...). The References
      heading is left without a chapter number.
    - The stray Normal-styled "3. The Emergence of Voice AI and
      LLM Agents" line in Chapter 1 is converted to a proper
      Heading 2 with the numeric prefix removed (the other Heading 2
      titles in this chapter are unnumbered).
    - The Chapter 4 title "Case Study Description - Better.com..."
      uses an em-dash that displays inconsistently across renderers;
      normalize it to a colon.
    """
    chapter_targets = [
        ("Introduction", "1. Introduction"),
        ("Literature review", "2. Literature Review"),
        ("Literature Review", "2. Literature Review"),
        ("Research Methodology", "3. Research Methodology"),
        # Use a flexible match for the Chapter 4 title because the
        # source text uses a hyphen and an em-dash variant has been
        # seen in some renderers.
        ("Case Study Description", None),
    ]
    chapter4_replacement = (
        "4. Case Study Description: Better.com and the Deployment of "
        "Voice AI in Mortgage Origination"
    )

    for para in doc.paragraphs:
        if para.style.name != "Heading 1":
            continue
        text = para.text.strip()
        for needle, replacement in chapter_targets:
            if text == needle and replacement is not None:
                set_paragraph_text(para, replacement)
                break
            if needle == "Case Study Description" and text.startswith(
                "Case Study Description"
            ):
                set_paragraph_text(para, chapter4_replacement)
                break

    # Promote the misformatted "3. The Emergence of Voice AI and LLM
    # Agents" line (currently a Normal-styled paragraph in Chapter 1)
    # to a proper Heading 2.
    for para in doc.paragraphs:
        if para.style.name != "Normal":
            continue
        text = _normalize_ws(para.text).strip()
        if text in {
            "3. The Emergence of Voice AI and LLM Agents",
            "The Emergence of Voice AI and LLM Agents",
        }:
            set_paragraph_text(para, "The Emergence of Voice AI and LLM Agents")
            try:
                para.style = doc.styles["Heading 2"]
            except KeyError:
                pass
            break


def insert_operationalization_table(doc) -> None:
    """Insert the 3.7 operationalization table after the rewritten P135."""
    anchor = find_paragraph_by_prefix(doc, "A central methodological step is the operationalization")
    if anchor is None:
        return
    insert_table_after(
        anchor,
        headers=[
            "Business value dimension",
            "Indicators",
            "Data source",
            "Reliability",
        ],
        rows=[
            [
                "Operational efficiency",
                "Expense Ratio; Cost-to-Originate (CTO) estimate; expense growth",
                "Better and Rocket 10-K filings; MBA Quarterly Performance Reports",
                "High (audited financials); Medium (CTO estimate)",
            ],
            [
                "Productivity and scalability",
                "Revenue per Employee; headcount vs revenue; Betsy call volume",
                "Better 10-K and earnings releases; ElevenLabs case study",
                "High (audited); Medium (management-reported volumes)",
            ],
            [
                "Financial performance",
                "Total Net Revenue; Net Income / Net Loss; Profit Margin",
                "Better and Rocket 10-K filings; S-1 (2018-2021)",
                "High (audited financials)",
            ],
            [
                "Customer experience",
                "Management-reported NPS; Trustpilot star ratings; VADER "
                "sentiment; BERTopic topic shares; keyword prevalence",
                "Earnings transcripts (NPS); Trustpilot reviews (NLP corpus)",
                "Medium (NPS, unaudited); High (NLP, but post-Betsy n is small)",
            ],
            [
                "Integration mechanism (qualitative)",
                "Tinman platform; Betsy as voice layer; ElevenLabs as voice "
                "infrastructure; human escalation model",
                "Better 10-K; press releases; ElevenLabs/BusinessWire materials",
                "High for Better disclosures; Medium for vendor materials",
            ],
        ],
    )


def main() -> int:
    repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    draft_path = os.path.join(repo_root, "04_Drafts", "Main_Draft.docx")
    if not os.path.exists(draft_path):
        print(f"Draft not found at {draft_path}", file=sys.stderr)
        return 1

    doc = docx.Document(draft_path)

    # 1) Whole-paragraph rewrites (most surgical changes happen here).
    for prefix, new_text in PARAGRAPH_REWRITES:
        para = find_paragraph_by_prefix(doc, prefix)
        if para is None:
            print(f"[warn] No paragraph starting with: {prefix[:60]!r}")
            continue
        set_paragraph_text(para, new_text)

    # 2) Cross-cutting substitutions on every paragraph.
    for para in doc.paragraphs:
        for old, new in SUBSTITUTIONS:
            if old in para.text:
                replace_in_runs(para, old, new)

    # 3) Normalize chapter headings (numbering, capitalization, level).
    normalize_chapter_headings(doc)

    # 4) Insert operationalization table after the rewritten 3.7 paragraph.
    insert_operationalization_table(doc)

    doc.save(draft_path)
    print(f"Saved updated draft to {draft_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
